import sys, os
sys.path.insert(0, os.path.dirname(__file__))

import io
from fastapi import APIRouter, UploadFile, File, HTTPException, Body
from fastapi.responses import Response
from typing import Optional

from llm_client import LLMClient

router = APIRouter()

llm = LLMClient()

SCRIPT_NAMES = {
    "devanagari": "Devanagari script (Hindi/Marathi/Sanskrit)",
    "latin":      "Latin/Roman script",
    "telugu":     "Telugu script",
    "tamil":      "Tamil script",
    "kannada":    "Kannada script",
    "bengali":    "Bengali script",
    "gujarati":   "Gujarati script",
    "gurmukhi":   "Gurmukhi/Punjabi script",
    "malayalam":  "Malayalam script",
    "odia":       "Odia script",
}

SCRIPT_FONTS = {
    "devanagari": "Anek Devanagari",
    "telugu":     "Noto Sans Telugu",
    "tamil":      "Noto Sans Tamil",
    "kannada":    "Noto Sans Kannada",
    "bengali":    "Noto Sans Bengali",
    "gujarati":   "Noto Sans Gujarati",
    "gurmukhi":   "Noto Sans Gurmukhi",
    "malayalam":  "Noto Sans Malayalam",
    "odia":       "Noto Sans Oriya",
}


def _transliterate_batch(texts: list, target_script: str, provider: str) -> dict:
    """Transliterate a batch of text segments. Returns {original: transliterated}."""
    script_name = SCRIPT_NAMES.get(target_script, target_script)
    numbered = "\n".join(f"{i+1}: {t}" for i, t in enumerate(texts))

    system = (
        f"You are an expert transliterator. Convert the pronunciation of words into {script_name}. "
        f"IMPORTANT: Transliteration means writing the SOUND of words in a new script — "
        f"do NOT translate the meaning. Output exactly the same number of numbered lines as input."
    )
    user = (
        f"Transliterate each segment into {script_name}.\n"
        f"Output EXACTLY {len(texts)} lines, each prefixed with its number:\n"
        f"1: <transliteration of segment 1>\n"
        f"2: <transliteration of segment 2>\n\n"
        f"Segments:\n{numbered}\n\nTransliterations:"
    )

    raw = llm.call(system, user, provider=provider)

    result = {}
    for line in raw.splitlines():
        line = line.strip()
        if ": " in line:
            num_str, _, trans = line.partition(": ")
            try:
                num = int(num_str.strip())
                if 1 <= num <= len(texts):
                    result[texts[num - 1]] = trans.strip()
            except ValueError:
                pass

    # Fallback: keep original for any missing segments
    for t in texts:
        if t not in result:
            result[t] = t

    return result


@router.get("/")
def root():
    return {"status": "running", "service": "ai-transliterate", "version": "1.0.0"}


@router.post("/transliterate-docx")
async def transliterate_docx(
    file: UploadFile = File(...),
    target_script: str = Body("devanagari"),
    llm_provider: str = Body("anthropic"),
):
    """Transliterate a DOCX file — convert script of each word without changing meaning."""
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")

    content = await file.read()

    from docx import Document
    from docx.oxml.ns import qn
    import tempfile

    doc = Document(io.BytesIO(content))

    # Collect unique non-empty paragraphs
    text_set = set()

    def collect(container):
        for para in container.paragraphs:
            if para.text.strip():
                text_set.add(para.text)
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    collect(cell)

    collect(doc)
    for section in doc.sections:
        for obj in [section.header, section.footer, section.first_page_header,
                    section.first_page_footer, section.even_page_header, section.even_page_footer]:
            if obj:
                collect(obj)

    blocks = list(text_set)
    if not blocks:
        out_buf = io.BytesIO()
        doc.save(out_buf)
        return Response(
            content=out_buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="transliterated.docx"'},
        )

    # Batch transliterate
    batch_size = 15
    transliterated_map = {}
    for i in range(0, len(blocks), batch_size):
        batch = blocks[i:i + batch_size]
        try:
            transliterated_map.update(_transliterate_batch(batch, target_script, llm_provider))
        except Exception as e:
            print(f"[Transliteration] Batch failed: {e} — keeping originals")
            for t in batch:
                transliterated_map[t] = t

    font_name = SCRIPT_FONTS.get(target_script)

    def apply(container):
        for para in container.paragraphs:
            if para.text not in transliterated_map:
                continue
            new_text = transliterated_map[para.text]
            if para.runs:
                text_runs = [r for r in para.runs if r.text.strip()]
                if text_runs:
                    text_runs[0].text = new_text
                    if font_name:
                        text_runs[0].font.name = font_name
                        rPr = text_runs[0]._element.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn("w:cs"), font_name)
                        rFonts.set(qn("w:hAnsi"), font_name)
                        rFonts.set(qn("w:ascii"), font_name)
                    for r in text_runs[1:]:
                        r.text = ""
            else:
                run = para.add_run(new_text)
                if font_name:
                    run.font.name = font_name
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    apply(cell)

    apply(doc)
    for section in doc.sections:
        for obj in [section.header, section.footer, section.first_page_header,
                    section.first_page_footer, section.even_page_header, section.even_page_footer]:
            if obj:
                apply(obj)

    out_buf = io.BytesIO()
    doc.save(out_buf)
    print(f"[Transliteration] Done — {len(transliterated_map)} segments → {target_script}")
    return Response(
        content=out_buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="transliterated.docx"'},
    )


@router.post("/transliterate-pdf")
async def transliterate_pdf(
    file: UploadFile = File(...),
    target_script: str = Body("devanagari"),
    llm_provider: str = Body("anthropic"),
):
    """Transliterate a PDF file — extract text, transliterate, return annotated PDF."""
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only .pdf files are supported.")

    content = await file.read()

    import fitz

    doc = fitz.open(stream=content, filetype="pdf")

    # Collect unique line texts
    line_texts = set()
    for page in doc:
        raw = page.get_text("dict", flags=0)
        for block in raw.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                text = "".join(s["text"] for s in line.get("spans", [])).strip()
                if text:
                    line_texts.add(text)

    all_lines = list(line_texts)
    transliterated_map = {}
    batch_size = 15
    for i in range(0, len(all_lines), batch_size):
        batch = all_lines[i:i + batch_size]
        try:
            transliterated_map.update(_transliterate_batch(batch, target_script, llm_provider))
        except Exception as e:
            print(f"[Transliteration-PDF] Batch failed: {e}")
            for t in batch:
                transliterated_map[t] = t

    # Apply redact + insert
    fonts_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "fonts")
    font_filename_map = {
        "devanagari": "AnekDevanagari-Regular.ttf",
        "telugu":     "NotoSansTelugu-Regular.ttf",
        "tamil":      "NotoSansTamil-Regular.ttf",
        "kannada":    "NotoSansKannada-Regular.ttf",
        "bengali":    "NotoSansBengali-Regular.ttf",
        "gujarati":   "NotoSansGujarati-Regular.ttf",
        "gurmukhi":   "NotoSansGurmukhi-Regular.ttf",
        "malayalam":  "NotoSansMalayalam-Regular.ttf",
        "odia":       "NotoSansOriya-Regular.ttf",
    }
    font_file = font_filename_map.get(target_script)
    font_path = None
    if font_file:
        candidate = os.path.join(fonts_dir, font_file)
        if os.path.exists(candidate):
            font_path = candidate

    archive = fitz.Archive()
    if font_path:
        with open(font_path, "rb") as f:
            font_bytes = f.read()
        font_basename = os.path.basename(font_path)
        archive.add((font_bytes, font_basename))
        font_name_css = "DocFont"
        font_css = f'@font-face {{ font-family: "{font_name_css}"; src: url("{font_basename}"); }}'
    else:
        font_name_css = "Helvetica, sans-serif"
        font_css = ""

    import html as _html

    for page in doc:
        raw = page.get_text("dict", flags=0)
        redactions = []
        for block in raw.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                text = "".join(s["text"] for s in spans).strip()
                if text not in transliterated_map or transliterated_map[text] == text:
                    continue
                first = spans[0]
                color_int = first.get("color", 0)
                redactions.append({
                    "bbox": fitz.Rect(line["bbox"]),
                    "fontsize": first.get("size", 11),
                    "color_hex": "#{:02x}{:02x}{:02x}".format(
                        (color_int >> 16) & 0xFF, (color_int >> 8) & 0xFF, color_int & 0xFF
                    ),
                    "text": transliterated_map[text],
                })

        if not redactions:
            continue

        for r in redactions:
            page.add_redact_annot(r["bbox"], fill=(1, 1, 1))
        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

        for r in redactions:
            adj_size = r["fontsize"] * 0.90
            html_str = (
                f'<span style="font-family: {font_name_css}; font-size: {adj_size}pt; '
                f'color: {r["color_hex"]}; line-height: 1.0;">'
                f'{_html.escape(r["text"])}</span>'
            )
            try:
                page.insert_htmlbox(r["bbox"], html_str, css=font_css, archive=archive, scale_low=0.75)
            except Exception as e:
                print(f"[Transliteration-PDF] insert_htmlbox failed: {e}")

    out_buf = io.BytesIO()
    doc.save(out_buf, garbage=4, deflate=True)
    doc.close()

    return Response(
        content=out_buf.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="transliterated.pdf"'},
    )
