import io


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from a PDF or DOCX file."""
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return _from_pdf(file_bytes)
    if fname.endswith(".docx"):
        return _from_docx(file_bytes)
    raise ValueError(f"Unsupported file type: {filename}. Only PDF and DOCX are supported.")


def _from_pdf(data: bytes) -> str:
    import fitz
    doc = fitz.open(stream=data, filetype="pdf")
    pages = [page.get_text() for page in doc]
    doc.close()
    return "\n\n".join(p for p in pages if p.strip())


def _from_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n\n".join(parts)
