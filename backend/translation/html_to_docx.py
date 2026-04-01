from docx import Document
from docx.shared import Pt, Inches
from bs4 import BeautifulSoup
import io
import base64
import re
from docx.oxml.ns import qn
from docx.oxml import parse_xml

class HTMLToDocx:
    def __init__(self):
        self.doc = Document()
        
    def convert(self, html_content):
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 1. Process Headers/Footers
        header_div = soup.find('div', class_='docx-header')
        footer_div = soup.find('div', class_='docx-footer')
        
        if self.doc.sections:
            section = self.doc.sections[0]
            if header_div:
                header = section.header
                # Use the first paragraph if it exists, otherwise add one
                p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                p.text = ""
                self._process_element(header_div, container=header)
            
            if footer_div:
                footer = section.footer
                p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                p.text = ""
                self._process_element(footer_div, container=footer)

        # 2. Process the main content container if it exists
        container = soup.find('div', class_='page') or soup.body or soup
        self._process_element(container)
                
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output

    def _process_element(self, parent, container=None):
        """Recursively process HTML elements and map them to docx components"""
        target = container if container is not None else self.doc
        
        for element in parent.contents:
            if element.name == None: # NavigableString
                continue
                
            # Skip header/footer divs as they are processed separately
            if element.name == 'div' and ('docx-header' in element.get('class', []) or 'docx-footer' in element.get('class', [])):
                continue

            if element.name.startswith('h'):
                level = int(element.name[1])
                target.add_heading(element.get_text().strip(), level=level)
            elif element.name == 'p':
                self._add_paragraph(element, container=target)
            elif element.name in ['ul', 'ol']:
                self._add_list(element, container=target)
            elif element.name == 'table':
                self._add_table(element, container=target)
            elif element.name == 'img':
                # Handle images outside paragraphs
                self._add_image(target, element)
            elif element.name == 'div':
                # Recursive call for nested divs
                self._process_element(element, container=target)

    def _add_paragraph(self, element, p=None, container=None):
        if p is None:
            target = container if container is not None else self.doc
            p = target.add_paragraph()
        
        # Handle alignment
        style = element.get('style', '')
        if 'text-align: center' in style:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'text-align: right' in style:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif 'text-align: justify' in style:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for child in element.children:
            self._process_inline(p, child)

    def _process_inline(self, p, child):
        if child.name == None: # Text node
            p.add_run(str(child))
        elif child.name in ['strong', 'b']:
            p.add_run(child.get_text()).bold = True
        elif child.name in ['em', 'i']:
            p.add_run(child.get_text()).italic = True
        elif child.name == 'u':
            p.add_run(child.get_text()).underline = True
        elif child.name == 'span':
            # Handle span with styles
            run = p.add_run(child.get_text())
            style = child.get('style', '')
            if 'font-weight: bold' in style: run.bold = True
            if 'font-style: italic' in style: run.italic = True
            if 'text-decoration: underline' in style: run.underline = True
        elif child.name == 'img':
            self._add_image(p, child)
        elif child.name == 'br':
            p.add_run('\n')
        else:
            # Fallback for unknown tags: process children or just text
            if hasattr(child, 'children'):
                for subchild in child.children:
                    self._process_inline(p, subchild)
            else:
                p.add_run(child.get_text())

    def _add_image(self, container, img_tag):
        src = img_tag.get('src', '')
        style = img_tag.get('style', '')
        
        # Parse width from style (e.g., width: 154.5pt)
        width_pt = None
        width_match = re.search(r'width:\s*([\d.]+)pt', style)
        if width_match:
            width_pt = float(width_match.group(1))

        if src.startswith('data:image/'):
            try:
                header, encoded = src.split(",", 1)
                image_data = base64.b64decode(encoded)
                image_stream = io.BytesIO(image_data)
                
                # Use captured width if available, otherwise fallback to default
                width_arg = Pt(width_pt) if width_pt else Inches(4)
                
                # If adding to a paragraph (run)
                if hasattr(container, 'add_run'):
                    run = container.add_run()
                    run.add_picture(image_stream, width=width_arg)
                else: # Adding directly to document/header/footer
                    container.add_picture(image_stream, width=width_arg)
            except Exception as e:
                print(f"Error restoring image in DOCX: {e}")

    def _add_list(self, element, container=None):
        target = container if container is not None else self.doc
        style = 'List Bullet' if element.name == 'ul' else 'List Number'
        for li in element.find_all('li', recursive=False):
            p = target.add_paragraph(style=style)
            for child in li.children:
                self._process_inline(p, child)

    def _add_table(self, element, container=None):
        target = container if container is not None else self.doc
        rows = element.find_all('tr', recursive=False)
        if not rows:
            return
            
        # Determine max columns
        num_cols = 0
        for row in rows:
            num_cols = max(num_cols, len(row.find_all(['td', 'th'], recursive=False)))
            
        if num_cols == 0:
            return

        table = target.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'], recursive=False)
            for j, cell in enumerate(cells):
                if j < num_cols:
                    docx_cell = table.cell(i, j)
                    
                    # Restore shading (background color)
                    cell_style = cell.get('style', '')
                    bg_match = re.search(r'background-color:\s*#([0-9a-fA-F]+)', cell_style)
                    if bg_match:
                        try:
                            shading_elm = parse_xml(f'<w:shd {qn("w:fill")}="{bg_match.group(1)}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                            docx_cell._tc.get_or_add_tcPr().append(shading_elm)
                        except: pass

                    # Clear default paragraph
                    docx_cell.text = ""
                    # Process content into cell
                    # Headers and Footers don't support tables directly in some contexts but cells do
                    for sub_el in cell.contents:
                        if sub_el.name == 'p':
                            self._add_paragraph(sub_el, p=docx_cell.add_paragraph())
                        elif sub_el.name == 'img':
                            p = docx_cell.add_paragraph()
                            self._add_image(p, sub_el)
                        elif sub_el.name == None:
                            if str(sub_el).strip():
                                docx_cell.add_paragraph(str(sub_el))
                        else:
                            p = docx_cell.add_paragraph()
                            self._process_inline(p, sub_el)

def html_to_docx(html_content):
    converter = HTMLToDocx()
    return converter.convert(html_content)
