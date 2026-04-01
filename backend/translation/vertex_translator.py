import os
import re
import time
from typing import Dict, Tuple
from docx import Document
from base_translator import BaseTranslator


class VertexTranslator(BaseTranslator):
    """Translates via Google Vertex AI (Gemini). Used as the 'gemini' provider in server flow."""

    PROTECTED_TERMS = [
        "L&T Finance Holdings Limited", "L&T Finance Holdings",
        "L&T Finance Limited", "L&T Housing Finance Limited",
        "L&T Housing Finance", "L&T Finance", "L&T",
        "Pvt. Ltd.", "Pvt Ltd", "Ltd.", "Co.", "Inc.", "Corp.", "Limited", "M/s",
        "NACH", "CIBIL", "NEFT", "RTGS", "MSME", "NBFC", "FOIR",
        "RBI", "SEBI", "GST", "PAN", "TDS", "EMI", "UPI", "KYC",
        "NPA", "MOU", "LOA", "NOC", "CIN", "DIN", "LLPIN", "SRN",
        "ROI", "IRR", "APR",
    ]

    LANGUAGE_NAMES = {
        'hi': 'Hindi', 'te': 'Telugu', 'mr': 'Marathi', 'kn': 'Kannada',
        'ta': 'Tamil', 'bn': 'Bengali', 'gu': 'Gujarati', 'or': 'Odia',
        'pa': 'Punjabi', 'as': 'Assamese', 'ml': 'Malayalam', 'rajasthani': 'Hindi',
    }

    def __init__(self):
        import vertexai
        project = os.getenv("VERTEX_PROJECT", "")
        location = os.getenv("VERTEX_LOCATION", "asia-south1")
        self.model_name = os.getenv("VERTEX_MODEL", "gemini-2.5-flash")
        self.max_tokens = 8192
        self.max_retries = 4

        if not project:
            raise ValueError("VERTEX_PROJECT must be set")

        vertexai.init(project=project, location=location)
        self.wingdings_map = {
            '\uF0A7': '☑', '\uF0A8': '☐',
            '\uF071': '✓', '\uF072': '✗',
            '\uF06F': '☐', '\uF0FE': '■',
        }
        print(f"[VertexTranslator] Initialized: model={self.model_name}, project={project}, location={location}")

    def _protect_terms(self, text: str) -> Tuple[str, Dict[str, str]]:
        term_map: Dict[str, str] = {}
        for i, term in enumerate(self.PROTECTED_TERMS):
            if term in text:
                token = f"__TERM_{i}__"
                term_map[token] = term
                text = text.replace(term, token)
        for j, bt in enumerate(re.findall(r'\[\[.*?\]\]', text)):
            token = f"__BRACKET_{j}__"
            if token not in term_map:
                term_map[token] = bt
                text = text.replace(bt, token, 1)
        return text, term_map

    def _restore_terms(self, text: str, term_map: Dict[str, str]) -> str:
        for token, original in term_map.items():
            text = text.replace(token, original)
        return text

    # ------------------------------------------------------------------
    # Internal helper
    # ------------------------------------------------------------------

    def _call_vertex(self, system_prompt: str, user_message: str) -> str:
        from vertexai.generative_models import GenerativeModel, GenerationConfig
        import google.api_core.exceptions

        model = GenerativeModel(self.model_name, system_instruction=system_prompt)

        for attempt in range(self.max_retries):
            try:
                response = model.generate_content(
                    user_message,
                    generation_config=GenerationConfig(
                        max_output_tokens=self.max_tokens,
                        temperature=0.0,
                    ),
                )
                return response.text
            except google.api_core.exceptions.ResourceExhausted:
                wait = 10 * (attempt + 1)  # 10s, 20s, 30s, 40s
                print(f"[Vertex] Rate limit, retrying in {wait}s (attempt {attempt+1}/{self.max_retries})")
                time.sleep(wait)
            except google.api_core.exceptions.GoogleAPIError:
                raise

        raise RuntimeError(f"[Vertex] API failed after {self.max_retries} retries")

    # ------------------------------------------------------------------
    # BaseTranslator interface
    # ------------------------------------------------------------------

    def translate_html(self, html_content: str, target_language: str) -> Dict:
        target_lang_name = self.LANGUAGE_NAMES.get(target_language, target_language)

        prompt = f"""You are an expert document translator specializing in corporate and regulatory documents. Translate the following HTML document to {target_lang_name}.

═══════════════════════════════════════════════════════════
SECTION A: HTML STRUCTURE PRESERVATION (ABSOLUTE RULES)
═══════════════════════════════════════════════════════════

1. **PRESERVE ALL HTML TAGS EXACTLY** - Every <div>, <p>, <span>, <table>, <tr>, <td>, <th>, <h1>-<h6>, <ul>, <ol>, <li>, <br>, <hr> must remain unchanged.
2. **PRESERVE ALL <style> BLOCKS** - Do NOT translate or modify anything inside <style>...</style> tags.
3. **PRESERVE ALL INLINE STYLES** - Every style="..." attribute must remain exactly as-is.
4. **PRESERVE ALL CLASS & ID ATTRIBUTES** - class="page", class="docx-header", class="docx-footer", etc. must NOT change.
5. **PRESERVE ALL TABLE STRUCTURE** - Maintain the EXACT number of <tr> and <td> tags. NEVER merge cells or skip rows.
6. **PRESERVE ALL IMAGES** - <img> tags must be kept exactly. DO NOT modify the `src="..."` attribute.
7. **PRESERVE [[IMG_PLACEHOLDER_N]] TOKENS** - Keep them exactly as-is (e.g., [[IMG_PLACEHOLDER_0]]).
8. **DO NOT ADD OR REMOVE ANY TAGS** - The output must have the identical tag structure as the input.

═══════════════════════════════════════════════════════════
SECTION B: WHAT TO TRANSLATE
═══════════════════════════════════════════════════════════

Translate ONLY the visible text content that appears BETWEEN HTML tags:
- Paragraph text, Headings, Table cell text, List items.
- Text content inside `docx-header` and `docx-footer` divs.

═══════════════════════════════════════════════════════════
SECTION C: DO NOT TRANSLATE
═══════════════════════════════════════════════════════════

- **Company Names:** L&T, L&T Finance, L&T Finance Limited, M/s, Pvt Ltd, Co., Ltd., Inc., Corp.
- **Financial Terms:** RBI, SEBI, GST, PAN, TDS, EMI, NACH, CIBIL, KYC, NEFT, RTGS, UPI, MSME, NBFC, etc.
- **Technical IDs:** Any string inside brackets like [[...]]

Return ONLY the translated HTML. No explanations, no markdown code blocks.

HTML to translate:
{html_content}

Translated HTML:"""

        system_prompt = (
            "You are an expert corporate document translator. Your ONLY job is to translate "
            "text content between HTML tags while preserving ALL HTML structure, CSS, styles, "
            "images, tables, and layout EXACTLY as given. You specialize in Indian regulatory "
            "and financial documents."
        )

        translated_html = self._call_vertex(system_prompt, prompt)

        if translated_html.startswith("```"):
            lines = translated_html.split('\n')
            if len(lines) > 2:
                translated_html = '\n'.join(lines[1:-1])

        return {"translated_html": translated_html, "language": target_language}

    def translate_docx(self, docx_path: str, target_language: str, output_path: str) -> str:
        doc = Document(docx_path)
        target_lang_name = self.LANGUAGE_NAMES.get(target_language.lower(), target_language)

        text_blocks = set()

        def collect_from_container(container):
            for para in container.paragraphs:
                if para.text.strip():
                    text = para.text
                    for char, mapped in self.wingdings_map.items():
                        text = text.replace(char, mapped)
                    text_blocks.add(text)
            for table in getattr(container, 'tables', []):
                for row in table.rows:
                    for cell in row.cells:
                        collect_from_container(cell)

        collect_from_container(doc)
        for section in doc.sections:
            for obj in [
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer,
            ]:
                if obj:
                    collect_from_container(obj)

        blocks = list(text_blocks)
        if not blocks:
            doc.save(output_path)
            return output_path

        batch_size = 25
        batches = [blocks[i:i + batch_size] for i in range(0, len(blocks), batch_size)]

        def _translate_batch(batch):
            try:
                protected_batch, batch_term_maps = [], []
                for seg in batch:
                    p, tm = self._protect_terms(seg)
                    protected_batch.append(p)
                    batch_term_maps.append(tm)

                prompt = f"""Translate each text segment to {target_lang_name}.

Return EXACTLY {len(batch)} lines — one translation per input segment.
- __TERM_N__ and __BRACKET_N__ tokens → copy verbatim, never translate.
- Checkboxes (☑ ☐ ✓ ✗ ■) → copy exactly as-is.
- No explanations, numbering, or markdown.

Segments:
---
{chr(10).join(protected_batch)}
---

Translated ({len(batch)} lines):"""

                system_prompt = (
                    f"Professional corporate translator. Output exactly {len(batch)} lines "
                    f"matching the input count. Never add commentary or formatting. "
                    f"__TERM_N__ and __BRACKET_N__ tokens must be copied verbatim."
                )

                raw_output = self._call_vertex(system_prompt, prompt)
                if raw_output.startswith("```"):
                    lines = raw_output.split('\n')
                    if len(lines) > 2:
                        raw_output = '\n'.join(lines[1:-1]).strip()
                batch_translated_raw = [t.strip() for t in raw_output.split('\n') if t.strip()]
                batch_translated = []
                for idx, translated_seg in enumerate(batch_translated_raw):
                    if idx < len(batch_term_maps):
                        translated_seg = self._restore_terms(translated_seg, batch_term_maps[idx])
                    batch_translated.append(translated_seg)
                return dict(zip(batch, batch_translated))
            except Exception as e:
                print(f"[Vertex] Batch translation failed, keeping originals: {e}")
                return {seg: seg for seg in batch}

        translated_map = {}
        for i, batch in enumerate(batches):
            translated_map.update(_translate_batch(batch))
            if i < len(batches) - 1:
                time.sleep(3)  # small gap between batches to respect quota

        from docx.oxml.ns import qn

        def apply_to_container(container):
            is_hindi = target_language.lower() in ['hi', 'rajasthani']
            for para in container.paragraphs:
                lookup_text = para.text
                for char, mapped in self.wingdings_map.items():
                    lookup_text = lookup_text.replace(char, mapped)
                if lookup_text not in translated_map:
                    continue
                translated_text = translated_map[lookup_text]
                if para.runs:
                    text_run_indices = [
                        i for i, run in enumerate(para.runs)
                        if not run._element.findall('.//' + qn('w:drawing'))
                        and not run._element.findall('.//' + qn('w:pict'))
                    ]
                    if text_run_indices:
                        main_run = para.runs[text_run_indices[0]]
                        main_run.text = translated_text
                        if is_hindi:
                            main_run.font.name = 'Anek Devanagari'
                            rPr = main_run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:cs'), 'Anek Devanagari')
                            rFonts.set(qn('w:hAnsi'), 'Anek Devanagari')
                            rFonts.set(qn('w:ascii'), 'Anek Devanagari')
                        for idx in text_run_indices[1:]:
                            para.runs[idx].text = ""
                    else:
                        new_run = para.add_run(translated_text)
                        if is_hindi:
                            new_run.font.name = 'Anek Devanagari'
                            rPr = new_run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:cs'), 'Anek Devanagari')
                else:
                    new_run = para.add_run(translated_text)
                    if is_hindi:
                        new_run.font.name = 'Anek Devanagari'
                        rPr = new_run._element.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:cs'), 'Anek Devanagari')

            for table in getattr(container, 'tables', []):
                for row in table.rows:
                    for cell in row.cells:
                        apply_to_container(cell)

        apply_to_container(doc)
        for section in doc.sections:
            for obj in [
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer,
            ]:
                if obj:
                    apply_to_container(obj)

        doc.save(output_path)
        return output_path
