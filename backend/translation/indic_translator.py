from docx import Document
import requests
import os
from typing import Dict, Optional
from base_translator import BaseTranslator

class IndicTransTranslator(BaseTranslator):
    def __init__(self):
        """Initialize IndicTrans2 client with API key and URL from environment"""
        self.api_key = os.getenv("INDIC_TRANS2_API_KEY")
        self.api_url = os.getenv("INDIC_TRANS2_API_URL")
    
    def translate_html(self, html_content: str, target_language: str) -> Dict:
        """
        Translate HTML content while preserving formatting using IndicTrans2
        
        Args:
            html_content: The HTML string to translate
            target_language: Target language code
        
        Returns:
            Dict with translated_html and language
        """
        if not self.api_url:
            raise ValueError("INDIC_TRANS2_API_URL environment variable not set")

        # Mapping for IndicTrans2 (standard codes)
        # Note: IndicTrans2 usually expects specific codes like 'hin_Deva' etc.
        # This implementation assumes the API handles standard 2-letter codes or is custom.
        
        headers = {
            "Authorization": f"Bearer {self.api_key}" if self.api_key else "",
            "Content-Type": "application/json"
        }
        
        data = {
            "text": html_content,
            "source_language": "eng_Latn",
            "target_language": target_language, # Should be mapping-aware in a real scenario
            "format": "html"
        }

        try:
            response = requests.post(self.api_url, json=data, headers=headers)
            response.raise_for_status()
            result = response.json()
            
            # Adjust based on actual API response format
            translated_html = result.get("translated_text", "")
            
            return {
                "translated_html": translated_html,
                "language": target_language
            }
        except Exception as e:
            print(f"IndicTrans2 Translation Error: {e}")
            raise Exception(f"IndicTrans2 translation failed: {str(e)}")
    def translate_docx(self, docx_path: str, target_language: str, output_path: str) -> str:
        """Translate DOCX natively using IndicTrans2 API"""
        doc = Document(docx_path)
        blocks = []
        for para in doc.paragraphs:
            if para.text.strip():
                blocks.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        blocks.append(cell.text)

        target_lang = self._map_lang(target_language)
        batch_size = 10
        translated_map = {}
        
        for i in range(0, len(blocks), batch_size):
            batch = blocks[i:i+batch_size]
            payload = {
                "inputs": batch,
                "source_language": "eng_Latn",
                "target_language": target_lang,
            }
            try:
                response = requests.post(self.api_url, json=payload, 
                                         headers={"Authorization": f"Bearer {self.api_key}" if self.api_key else ""})
                if response.ok:
                    batch_translated = response.json().get("outputs", [])
                    for original, translated in zip(batch, batch_translated):
                        translated_map[original] = translated
            except Exception as e:
                print(f"IndicTrans Batch Error: {e}")

        for para in doc.paragraphs:
            if para.text in translated_map:
                translated_text = translated_map[para.text]
                if para.runs:
                    first_run = para.runs[0]
                    first_run.text = translated_text
                    for i in range(1, len(para.runs)):
                        para.runs[i].text = ""
                else:
                    para.text = translated_text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text in translated_map:
                        translated_text = translated_map[cell.text]
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            first_run = cell.paragraphs[0].runs[0]
                            first_run.text = translated_text
                            for i in range(1, len(cell.paragraphs[0].runs)):
                                cell.paragraphs[0].runs[i].text = ""
                        else:
                            cell.text = translated_text

        doc.save(output_path)
        return output_path

    def _map_lang(self, lang_code):
        mapping = {
            "hi": "hin_Deva", "te": "tel_Telu", "mr": "mar_Deva",
            "bn": "ben_Beng", "kn": "kan_Knda", "ta": "tam_Taml",
            "gu": "guj_Gujr", "or": "ory_Orya", "ml": "mal_Mlym",
            "pa": "pan_Guru", "as": "asm_Asst"
        }
        return mapping.get(lang_code, "hin_Deva")
