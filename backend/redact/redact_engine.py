"""Core redaction helpers for PDF (PyMuPDF) and DOCX (python-docx) documents."""

import io
import re
from typing import Dict, List, Tuple


# ── Color utilities ────────────────────────────────────────────────────────────

_HEX_TO_WD_COLOR = {
    "#ffff00": "YELLOW",
    "#00ff00": "GREEN",
    "#ff00ff": "MAGENTA",
    "#ff0000": "RED",
    "#0000ff": "BLUE",
    "#00ffff": "CYAN",
    "#000000": "BLACK",
    "#ffffff": "WHITE",
    "#00008b": "DARK_BLUE",
    "#008b8b": "DARK_CYAN",
    "#006400": "DARK_GREEN",
    "#8b008b": "DARK_MAGENTA",
    "#8b0000": "DARK_RED",
    "#808000": "DARK_YELLOW",
    "#808080": "DARK_GRAY",
    "#d3d3d3": "LIGHT_GRAY",
}


def _hex_to_rgb_float(hex_color: str) -> Tuple[float, float, float]:
    h = hex_color.lstrip("#")
    if len(h) == 3:
        h = h[0] * 2 + h[1] * 2 + h[2] * 2
    return (int(h[0:2], 16) / 255.0, int(h[2:4], 16) / 255.0, int(h[4:6], 16) / 255.0)


def _colors_match(c1: tuple, c2: tuple, tol: float = 0.12) -> bool:
    return all(abs(a - b) <= tol for a, b in zip(c1, c2))


def _closest_wd_color(hex_color: str) -> str:
    """Return the WD_COLOR_INDEX attribute name closest to the given hex color."""
    target = _hex_to_rgb_float(hex_color)
    best, best_dist = "YELLOW", float("inf")
    for h, name in _HEX_TO_WD_COLOR.items():
        c = _hex_to_rgb_float(h)
        d = sum((a - b) ** 2 for a, b in zip(target, c))
        if d < best_dist:
            best_dist = d
            best = name
    return best


# ── PDF helpers ────────────────────────────────────────────────────────────────

def redact_pdf_by_phrases(pdf_bytes: bytes, phrases: List[str]) -> bytes:
    """Find every occurrence of each phrase in the PDF and apply a black redaction box."""
    import fitz

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        for phrase in phrases:
            phrase = phrase.strip()
            if not phrase:
                continue
            for rect in page.search_for(phrase):
                page.add_redact_annot(rect, fill=(0, 0, 0))
        page.apply_redactions()

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


def redact_pdf_by_highlight_color(pdf_bytes: bytes, hex_color: str) -> bytes:
    """Redact all text in highlight annotations matching the given hex color in a PDF."""
    import fitz

    r, g, b = _hex_to_rgb_float(hex_color)
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    for page in doc:
        to_delete = []
        for annot in page.annots():
            if annot.type[0] != 8:  # 8 = PDF_ANNOT_HIGHLIGHT
                continue
            colors = annot.colors
            stroke = colors.get("stroke") or colors.get("fill") or ()
            if stroke and _colors_match(stroke, (r, g, b)):
                page.add_redact_annot(annot.rect, fill=(0, 0, 0))
                to_delete.append(annot)
        for a in to_delete:
            page.delete_annot(a)
        page.apply_redactions()

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


def redact_pdf_by_positions(pdf_bytes: bytes, specs: List[Dict]) -> bytes:
    """
    Redact PDF lines by position.

    Each spec:
      { "page": int (1-indexed),
        "lines": [int, ...] (1-indexed line numbers on that page),
        "paragraph": int (optional, 1-indexed block on the page) }

    If "paragraph" is given, lines are counted within that block only.
    If "paragraph" is absent, lines are counted flat across all blocks on the page.
    If "lines" is absent/empty, the entire paragraph/block is redacted.
    """
    import fitz

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    for spec in specs:
        page_num = spec.get("page", 1) - 1  # 0-indexed
        line_nums = set(spec.get("lines") or [])
        para_num = spec.get("paragraph")

        if page_num < 0 or page_num >= doc.page_count:
            continue

        page = doc[page_num]
        page_dict = page.get_text("dict")
        blocks = [b for b in page_dict.get("blocks", []) if b.get("type") == 0]

        if para_num is not None:
            para_idx = para_num - 1
            if para_idx >= len(blocks):
                continue
            lines = blocks[para_idx].get("lines", [])
            for li, line in enumerate(lines, start=1):
                if not line_nums or li in line_nums:
                    page.add_redact_annot(fitz.Rect(line["bbox"]), fill=(0, 0, 0))
        else:
            all_lines = [ln for b in blocks for ln in b.get("lines", [])]
            for li, line in enumerate(all_lines, start=1):
                if li in line_nums:
                    page.add_redact_annot(fitz.Rect(line["bbox"]), fill=(0, 0, 0))

        page.apply_redactions()

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


# ── DOCX helpers ───────────────────────────────────────────────────────────────

def _redact_phrases_in_paragraph(para, phrases: List[str]) -> None:
    """
    Redact phrases within a single paragraph.
    - Tries single-run replacement first (handles most cases).
    - Falls back to whole-paragraph redaction when a phrase spans multiple runs.
    """
    full_lower = para.text.lower()

    for phrase in phrases:
        phrase = phrase.strip()
        if not phrase or phrase.lower() not in full_lower:
            continue

        replaced = False
        for run in para.runs:
            if phrase.lower() in run.text.lower():
                run.text = re.sub(re.escape(phrase), "[REDACTED]", run.text, flags=re.IGNORECASE)
                replaced = True

        if not replaced:
            # Phrase spans multiple runs — redact the whole paragraph safely
            for run in para.runs:
                if run.text.strip():
                    run.text = "[REDACTED]"
            break


def redact_docx_by_phrases(docx_bytes: bytes, phrases: List[str]) -> bytes:
    """Replace every run containing a redact phrase with [REDACTED]."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))

    for para in doc.paragraphs:
        _redact_phrases_in_paragraph(para, phrases)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _redact_phrases_in_paragraph(para, phrases)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def redact_docx_by_highlight_color(docx_bytes: bytes, hex_color: str) -> bytes:
    """Redact runs highlighted with the given hex color in a DOCX file."""
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    normalized = hex_color.lower()
    if not normalized.startswith("#"):
        normalized = f"#{normalized}"

    color_name = _HEX_TO_WD_COLOR.get(normalized) or _closest_wd_color(normalized)
    target_idx = getattr(WD_COLOR_INDEX, color_name, None)

    doc = Document(io.BytesIO(docx_bytes))

    def _check_para(para):
        for run in para.runs:
            if target_idx is not None and run.font.highlight_color == target_idx:
                run.text = "[REDACTED]"
                run.font.highlight_color = None

    for para in doc.paragraphs:
        _check_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _check_para(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def redact_docx_by_positions(docx_bytes: bytes, specs: List[Dict]) -> bytes:
    """
    Redact DOCX by paragraph index (1-indexed from the whole document).

    Each spec: { "paragraph": int }

    Note: DOCX paragraphs don't have intrinsic line boundaries, so the entire
    paragraph is redacted when targeted.
    """
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    all_paras = doc.paragraphs
    para_indices = {spec["paragraph"] - 1 for spec in specs if "paragraph" in spec}

    for idx in para_indices:
        if 0 <= idx < len(all_paras):
            for run in all_paras[idx].runs:
                if run.text.strip():
                    run.text = "[REDACTED]"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
