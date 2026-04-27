"""
AWS Bedrock service for role-based document redaction.
Uses Claude vision via the Bedrock converse API — sends PDFs/images directly
without needing a separate text-extraction step.
"""

import asyncio
import functools
import io
import json
import logging
import os
import re
import time
from typing import List

import boto3
import botocore.config
import botocore.exceptions

logger = logging.getLogger(__name__)

# ── MIME helpers ───────────────────────────────────────────────────────────────

DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/docx",
    "application/zip",
}

# ── Role context prompts ───────────────────────────────────────────────────────

ROLE_CONTEXTS = {
    "legal_ops": (
        "The viewer is in Legal Operations. "
        "Redact: privileged attorney-client communications, internal counsel notes, litigation strategy, "
        "settlement amounts, and any text marked PRIVILEGED, CONFIDENTIAL — ATTORNEY CLIENT, or similar."
    ),
    "admin": (
        "The viewer is a System Administrator. "
        "Redact: personally identifiable information (PII), passwords, API keys, access credentials, "
        "social security numbers, bank account numbers, credit card numbers, and private employee records."
    ),
    "hr": (
        "The viewer is in Human Resources. "
        "Redact: salary and compensation figures, performance review details, disciplinary records, "
        "personal health information, background check results, and sensitive personal details."
    ),
    "finance": (
        "The viewer is in Finance. "
        "Redact: unreleased revenue projections, confidential pricing, M&A deal terms, "
        "loan account numbers, investor names, and non-public financial forecasts."
    ),
    "executive": (
        "The viewer is an Executive / C-Suite member. "
        "Redact only the most sensitive operational details: individual employee PII, "
        "raw customer data, and privileged legal opinions."
    ),
    "external": (
        "The viewer is an External Party (outside the organisation). Apply maximum redaction: "
        "redact all internal financial data, employee names and roles, internal project names, "
        "confidential business terms, pricing, strategy, legal matters, and any internal identifiers."
    ),
}

_REDACT_SUFFIX = (
    "\n\nIdentify ALL exact text phrases from the document that must be REDACTED for this viewer.\n\n"
    "Rules:\n"
    "1. Return ONLY valid JSON — no markdown fences, no commentary.\n"
    "2. Each entry must be VERBATIM text from the document.\n"
    "3. Be thorough — include names, numbers, account details, confidential figures.\n"
    "4. Short fragments (a single number, a name) are valid entries.\n\n"
    'Return: {"redact_phrases": ["exact phrase 1", "exact phrase 2", ...]}'
)

_SYSTEM = "You are a document security specialist. Identify sensitive text to redact based on the viewer's role."


# ── Shared helpers ─────────────────────────────────────────────────────────────

def _extract_docx_text(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n\n".join(parts)


def _pdf_to_images(file_bytes: bytes, max_pages: int = 8, dpi: int = 150) -> List[bytes]:
    """Render PDF pages to JPEG images using PyMuPDF."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        zoom = dpi / 72
        mat = fitz.Matrix(zoom, zoom)
        images = []
        for i in range(min(len(doc), max_pages)):
            pix = doc.load_page(i).get_pixmap(matrix=mat, alpha=False)
            images.append(pix.tobytes("jpeg"))
        doc.close()
        return images
    except Exception as exc:
        logger.warning("[BedrockRedactService] pdf_to_images failed: %s", exc)
        return []


def _extract_pdf_text(file_bytes: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "\n\n".join(page.get_text() for page in doc)
        doc.close()
        return text
    except Exception:
        return ""


def _parse_phrases(raw: str) -> List[str]:
    """Extract the redact_phrases list from LLM JSON output."""
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-z]*\n?", "", raw).rstrip("`").strip()
    try:
        return json.loads(raw).get("redact_phrases", [])
    except json.JSONDecodeError:
        pass
    # Try to find the array inline
    match = re.search(r'"redact_phrases"\s*:\s*(\[.*?\])', raw, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except Exception:
            pass
    logger.warning("[BedrockRedactService] could not parse phrases from: %s", raw[:300])
    return []


# ── Main service ───────────────────────────────────────────────────────────────

class BedrockRedactService:
    """
    Role-based redaction phrase identification via AWS Bedrock (Claude, vision-capable).

    Uses the Bedrock converse API:
      • PDF  → native document block (best quality), falls back to JPEG image pages
      • DOCX → text extraction then text prompt
      • Image → inline image block
    """

    def __init__(self):
        region = os.getenv("AWS_DEFAULT_REGION", "ap-south-1")
        self.model_id = os.getenv("BEDROCK_MODEL", "global.anthropic.claude-haiku-4-5-20251001-v1:0")
        self.max_tokens = int(os.getenv("MAX_OUTPUT_TOKENS", "8096"))
        self.max_retries = 3
        self.client = boto3.client(
            "bedrock-runtime",
            region_name=region,
            config=botocore.config.Config(
                read_timeout=300,
                connect_timeout=10,
                retries={"max_attempts": 0},
            ),
        )
        print(f"[BedrockRedactService] Initialized: model={self.model_id}, region={region}")
        logger.info("[BedrockRedactService] Initialized: model=%s, region=%s", self.model_id, region)

    # ── Internal sync helpers ──────────────────────────────────────────────────

    def _call_bedrock(self, content: list) -> str:
        for attempt in range(self.max_retries):
            try:
                response = self.client.converse(
                    modelId=self.model_id,
                    system=[{"text": _SYSTEM}],
                    messages=[{"role": "user", "content": content}],
                    inferenceConfig={"maxTokens": self.max_tokens, "temperature": 0.0},
                )
                if response.get("stopReason") == "max_tokens":
                    logger.warning("[BedrockRedactService] Hit max_tokens — output may be truncated")
                return response["output"]["message"]["content"][0]["text"]
            except botocore.exceptions.ClientError as e:
                code = e.response["Error"]["Code"]
                if code in ("ThrottlingException", "ServiceUnavailableException") and attempt < self.max_retries - 1:
                    wait = 2 ** attempt
                    logger.warning("[BedrockRedactService] %s — retrying in %ds (attempt %d/%d)", code, wait, attempt + 1, self.max_retries)
                    time.sleep(wait)
                else:
                    raise
        raise RuntimeError(f"[BedrockRedactService] Bedrock API failed after {self.max_retries} retries")

    def _identify_sync(self, file_bytes: bytes, mime_type: str, filename: str, role: str) -> List[str]:
        ext = (filename or "").lower().rsplit(".", 1)[-1]
        is_docx = mime_type in DOCX_MIMES or ext in ("docx", "doc")
        is_image = not is_docx and mime_type.startswith("image/")
        file_type = "DOCX" if is_docx else ("IMAGE" if is_image else "PDF")

        role_ctx = ROLE_CONTEXTS.get(role, f"The viewer's role is '{role}'.")
        prompt = f"{role_ctx}{_REDACT_SUFFIX}"

        print(f"[BedrockRedactService] identify | file={filename} | type={file_type} | role={role} | model={self.model_id}")

        # ── DOCX: text extraction path ─────────────────────────────────────────
        if is_docx:
            text = _extract_docx_text(file_bytes)
            if not text.strip():
                print(f"[BedrockRedactService] DOCX had no extractable text — returning empty")
                return []
            content = [{"text": f"Document text:\n\n{text[:12000]}\n\n---\n{prompt}"}]
            raw = self._call_bedrock(content)
            phrases = _parse_phrases(raw)
            print(f"[BedrockRedactService] DONE | phrases={len(phrases)}")
            return phrases

        # ── Image: inline image block ──────────────────────────────────────────
        if is_image:
            fmt = mime_type.split("/")[1].lower()
            if fmt not in ("jpeg", "png", "gif", "webp"):
                fmt = "jpeg"
            content = [
                {"image": {"format": fmt, "source": {"bytes": file_bytes}}},
                {"text": prompt},
            ]
            raw = self._call_bedrock(content)
            phrases = _parse_phrases(raw)
            print(f"[BedrockRedactService] DONE | phrases={len(phrases)}")
            return phrases

        # ── PDF: native document block → image pages → text fallback ──────────
        try:
            content = [
                {"document": {"format": "pdf", "name": "document", "source": {"bytes": file_bytes}}},
                {"text": prompt},
            ]
            raw = self._call_bedrock(content)
            phrases = _parse_phrases(raw)
            print(f"[BedrockRedactService] DONE (pdf-native) | phrases={len(phrases)}")
            return phrases
        except Exception as exc:
            logger.warning("[BedrockRedactService] PDF document block failed (%s) — trying image pages", exc)

        page_images = _pdf_to_images(file_bytes)
        if page_images:
            content = [{"image": {"format": "jpeg", "source": {"bytes": img}}} for img in page_images]
            content.append({"text": prompt})
            raw = self._call_bedrock(content)
            phrases = _parse_phrases(raw)
            print(f"[BedrockRedactService] DONE (pdf-images) | phrases={len(phrases)}")
            return phrases

        # Last resort: plain text
        text = _extract_pdf_text(file_bytes)
        content = [{"text": f"Document text:\n\n{text[:12000]}\n\n---\n{prompt}"}]
        raw = self._call_bedrock(content)
        phrases = _parse_phrases(raw)
        print(f"[BedrockRedactService] DONE (pdf-text) | phrases={len(phrases)}")
        return phrases

    # ── Public async interface ─────────────────────────────────────────────────

    async def identify_phrases_for_role(
        self,
        file_bytes: bytes,
        mime_type: str,
        filename: str,
        role: str,
    ) -> List[str]:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            functools.partial(self._identify_sync, file_bytes, mime_type, filename, role),
        )
