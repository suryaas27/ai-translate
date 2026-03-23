import os
import time
from typing import Dict
from docx import Document
from openai import OpenAI
from base_translator import BaseTranslator


class SarvamTranslator(BaseTranslator):
    def __init__(self):
        api_key = os.getenv("SARVAM_API_KEY")
        if not api_key:
            raise ValueError("SARVAM_API_KEY environment variable not set")
        # Sarvam-M uses an OpenAI-compatible chat completions endpoint
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://api.sarvam.ai/v1",
        )

    def translate_html(self, html_content: str, target_language: str) -> Dict:
        language_names = {
            'hi': 'Hindi', 'te': 'Telugu', 'mr': 'Marathi',
            'bn': 'Bengali', 'kn': 'Kannada', 'ta': 'Tamil',
            'gu': 'Gujarati', 'or': 'Odia', 'ml': 'Malayalam',
            'pa': 'Punjabi', 'as': 'Assamese', 'rajasthani': 'Hindi'
        }
        target_lang_name = language_names.get(target_language, target_language)

        prompt = f"""You are an expert document translator specializing in corporate and regulatory documents. Translate the following HTML document to {target_lang_name}.

═══════════════════════════════════════════════════════════
SECTION A: HTML STRUCTURE PRESERVATION (ABSOLUTE RULES)
═══════════════════════════════════════════════════════════

1. **PRESERVE ALL HTML TAGS EXACTLY** - Every <div>, <p>, <span>, <table>, <tr>, <td>, <th>, <h1>-<h6>, <ul>, <ol>, <li>, <br>, <hr> must remain unchanged.
2. **PRESERVE ALL <style> BLOCKS** - Do NOT translate or modify anything inside <style>...</style> tags.
3. **PRESERVE ALL INLINE STYLES** - Every style="..." attribute must remain exactly as-is.
4. **PRESERVE ALL CLASS & ID ATTRIBUTES** - class="page", class="docx-header", etc. must NOT change.
5. **PRESERVE ALL TABLE STRUCTURE** - Maintain the EXACT number of <tr> and <td> tags. NEVER merge cells or skip rows.
6. **PRESERVE ALL IMAGES** - <img> tags must be kept exactly. DO NOT modify src="..." attributes.
7. **PRESERVE [[IMG_PLACEHOLDER_N]] TOKENS** - Keep them exactly as-is (e.g., [[IMG_PLACEHOLDER_0]]).
8. **DO NOT ADD OR REMOVE ANY TAGS** - The output must have identical tag structure as the input.

═══════════════════════════════════════════════════════════
SECTION B: WHAT TO TRANSLATE
═══════════════════════════════════════════════════════════

Translate ONLY the visible text content that appears BETWEEN HTML tags:
- Paragraph text, Headings, Table cell text, List items.
- Text content inside `docx-header` and `docx-footer` divs.

═══════════════════════════════════════════════════════════
SECTION C: DO NOT TRANSLATE
═══════════════════════════════════════════════════════════

- **Company Names:** L&T, L&T Finance, L&T Finance Limited, M/s, Pvt Ltd, Co., Ltd., Inc., Corp.
- **Financial Terms:** RBI, SEBI, GST, PAN, TDS, EMI, NACH, CIBIL, KYC, NEFT, RTGS, UPI, MSME
- **Technical IDs:** Any string inside brackets like [[...]]

Return ONLY the translated HTML. No explanations, no markdown code blocks.

HTML to translate:
{html_content}

Translated HTML:"""

        for attempt in range(4):
            try:
                response = self.client.chat.completions.create(
                    model="sarvam-m",
                    messages=[{"role": "user", "content": prompt}],
                )
                break
            except Exception as e:
                if '429' in str(e) or 'rate_limit' in str(e).lower():
                    wait = 10 * (attempt + 1)
                    print(f"[Sarvam] Rate limit, retrying in {wait}s (attempt {attempt+1}/4)")
                    time.sleep(wait)
                else:
                    raise
        else:
            raise RuntimeError("[Sarvam] Rate limit: all retries exhausted")

        translated_html = response.choices[0].message.content.strip()
        if translated_html.startswith("```"):
            lines = translated_html.split('\n')
            if len(lines) > 2:
                translated_html = '\n'.join(lines[1:-1])

        return {"translated_html": translated_html, "language": target_language}

    def translate_docx(self, docx_path: str, target_language: str, output_path: str) -> str:
        """Translate DOCX natively using sarvam-m with full layout preservation."""
        from docx.oxml.ns import qn
        doc = Document(docx_path)

        language_names = {
            'hi': 'Hindi', 'te': 'Telugu', 'mr': 'Marathi', 'kn': 'Kannada',
            'ta': 'Tamil', 'bn': 'Bengali', 'gu': 'Gujarati', 'or': 'Odia',
            'pa': 'Punjabi', 'as': 'Assamese', 'ml': 'Malayalam', 'rajasthani': 'Hindi'
        }
        target_lang_name = language_names.get(target_language.lower(), target_language)

        text_blocks = set()

        def collect_from_container(container):
            for para in container.paragraphs:
                if para.text.strip():
                    text_blocks.add(para.text)
            for table in getattr(container, 'tables', []):
                for row in table.rows:
                    for cell in row.cells:
                        collect_from_container(cell)

        collect_from_container(doc)
        for section in doc.sections:
            for obj in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
                if obj:
                    collect_from_container(obj)

        blocks = list(text_blocks)
        if not blocks:
            doc.save(output_path)
            return output_path

        batch_size = 25
        batches = [blocks[i:i+batch_size] for i in range(0, len(blocks), batch_size)]

        def _translate_batch(batch):
            prompt = f"""You are a professional document translator specializing in corporate legal and regulatory documents.
Translate the following Word document segments to {target_lang_name}.

CRITICAL RULES:
1. Translate ONLY the text content provided.
2. Return precisely the same number of segments as input.
3. Separate each translated segment with a single newline (\\n).
4. DO NOT add any explanations, numeric indices, or formatting.
5. PRESERVE company names exactly: "L&T", "L&T Finance", "M/s", "Pvt Ltd", "Ltd.", "Co.", "Limited".
6. PRESERVE numbers, dates, and technical abbreviations (RBI, SEBI, GST, PAN, TDS, EMI, NACH, CIBIL, KYC).
7. PRESERVE checkboxes (☑, ☐) and symbols (✓, ✗) exactly as they appear.

Segments to translate:
---
{chr(10).join(batch)}
---

Translated segments:"""
            for attempt in range(4):
                try:
                    response = self.client.chat.completions.create(
                        model="sarvam-m",
                        messages=[{"role": "user", "content": prompt}],
                    )
                    break
                except Exception as e:
                    if '429' in str(e) or 'rate_limit' in str(e).lower():
                        wait = 10 * (attempt + 1)
                        print(f"[Sarvam] Rate limit in batch, retrying in {wait}s (attempt {attempt+1}/4)")
                        time.sleep(wait)
                    else:
                        raise
            else:
                raise RuntimeError("[Sarvam] Rate limit: all retries exhausted")

            raw = response.choices[0].message.content.strip()
            if raw.startswith("```"):
                lines = raw.split('\n')
                if len(lines) > 2:
                    raw = '\n'.join(lines[1:-1]).strip()
            translated = [t.strip() for t in raw.split('\n') if t.strip()]
            return dict(zip(batch, translated))

        translated_map = {}
        for i, batch in enumerate(batches):
            translated_map.update(_translate_batch(batch))
            if i < len(batches) - 1:
                time.sleep(3)

        is_hindi = target_language.lower() in ['hi', 'rajasthani']

        def apply_to_container(container):
            for para in container.paragraphs:
                if para.text in translated_map:
                    translated_text = translated_map[para.text]
                    if para.runs:
                        text_run_indices = [
                            i for i, run in enumerate(para.runs)
                            if not run._element.findall('.//' + qn('w:drawing'))
                            and not run._element.findall('.//' + qn('w:pict'))
                        ]
                        if text_run_indices:
                            main_run = para.runs[text_run_indices[0]]
                            main_run.text = translated_text
                            if is_hindi:
                                main_run.font.name = 'Anek Devanagari'
                                rPr = main_run._element.get_or_add_rPr()
                                rFonts = rPr.get_or_add_rFonts()
                                rFonts.set(qn('w:cs'), 'Anek Devanagari')
                                rFonts.set(qn('w:hAnsi'), 'Anek Devanagari')
                                rFonts.set(qn('w:ascii'), 'Anek Devanagari')
                            for idx in text_run_indices[1:]:
                                para.runs[idx].text = ""
                        else:
                            new_run = para.add_run(translated_text)
                            if is_hindi:
                                new_run.font.name = 'Anek Devanagari'
                    else:
                        new_run = para.add_run(translated_text)
                        if is_hindi:
                            new_run.font.name = 'Anek Devanagari'

            for table in getattr(container, 'tables', []):
                for row in table.rows:
                    for cell in row.cells:
                        apply_to_container(cell)

        apply_to_container(doc)
        for section in doc.sections:
            for obj in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
                if obj:
                    apply_to_container(obj)

        doc.save(output_path)
        return output_path
