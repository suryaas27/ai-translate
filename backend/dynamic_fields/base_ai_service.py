"""Abstract base class for AI services used in dynamic field extraction."""

import base64
import io
import json
import logging
import re
from abc import ABC, abstractmethod
from typing import Dict, List

logger = logging.getLogger(__name__)

# ── Shared prompt ──────────────────────────────────────────────────────────────

TEMPLATE_EXTRACTION_PROMPT = (
    "You are a document template analyzer. Your ONLY job is to find every blank that "
    "someone is expected to fill in.\n\n"
    "Blank fields look like ANY of these:\n"
    "  • Underscores   → Name: _____________  or  Name:___\n"
    "  • Dashes        → Date: -------------\n"
    "  • Dots          → Address: ...........\n"
    "  • Brackets      → [EMPLOYEE NAME]  or  <Company>\n"
    "  • Empty parens  → Signed: (             )\n"
    "  • Dollar tokens → $name  or  $Date  or  $CompanyName  (a $ followed by a word)\n"
    "  • Visible underlines (a line drawn under empty space in a form)\n\n"
    "Rules:\n"
    "1. Replace EVERY blank with a unique marker {{field_0}}, {{field_1}}, … "
    "(keep incrementing even if blanks look the same).\n"
    "2. Copy ALL other text exactly — do NOT summarise or skip content.\n"
    "3. Infer a short label from the nearest context "
    "(e.g. 'Name: ____' → 'Name', 'Date of Birth:----' → 'Date of Birth').\n"
    "4. If zero blanks are found, still return the full text and an empty fields list.\n\n"
    "Return ONLY valid JSON, no markdown fences, no explanation:\n"
    '{"fields":[{"id":"field_0","label":"Employee Name"},...],'
    '"template_text":"Full document text here with {{field_N}} in place of blanks."}'
)

# MIME types recognised as DOCX
DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/docx",
    "application/zip",
}


# ── Shared static helpers ──────────────────────────────────────────────────────

def pdf_to_images(file_bytes: bytes, max_pages: int = 5, dpi: int = 150) -> List[bytes]:
    """Render PDF pages to JPEG images using PyMuPDF."""
    try:
        import fitz
    except ImportError:
        return []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        images: List[bytes] = []
        zoom = dpi / 72
        mat = fitz.Matrix(zoom, zoom)
        for page_idx in range(min(len(doc), max_pages)):
            page = doc.load_page(page_idx)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            images.append(pix.tobytes("jpeg"))
        doc.close()
        return images
    except Exception as exc:
        logger.warning("pdf_to_images failed: %s", exc)
        return []


def extract_pdf_text_naive(file_bytes: bytes) -> str:
    """Extract text from a PDF using pypdf — pure Python fallback."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        logger.warning("extract_pdf_text_naive failed: %s", exc)
        return ""


def extract_docx_text(file_bytes: bytes) -> str:
    """Return all paragraph and table-cell text from a DOCX file."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines: List[str] = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        lines.append(para.text)
    return "\n".join(lines)


def parse_template_json(raw: str, source: str) -> Dict:
    """Parse LLM JSON output with truncation recovery."""
    raw = re.sub(r"```(?:json)?", "", raw).strip().rstrip("`")

    try:
        parsed = json.loads(raw)
        return {
            "fields": parsed.get("fields", []),
            "template_text": parsed.get("template_text", ""),
        }
    except json.JSONDecodeError:
        pass

    # Salvage truncated output
    fields_match = re.search(r'"fields"\s*:\s*\[', raw)
    if fields_match:
        array_start = fields_match.end() - 1
        salvaged: list = []
        i = array_start + 1
        while i < len(raw):
            while i < len(raw) and raw[i] in (' ', '\n', '\r', '\t', ','):
                i += 1
            if i >= len(raw) or raw[i] != '{':
                break
            depth = 0
            j = i
            while j < len(raw):
                if raw[j] == '{':
                    depth += 1
                elif raw[j] == '}':
                    depth -= 1
                    if depth == 0:
                        break
                j += 1
            try:
                obj = json.loads(raw[i:j + 1])
                salvaged.append(obj)
            except json.JSONDecodeError:
                break
            i = j + 1
        if salvaged:
            tmpl_match = re.search(r'"template_text"\s*:\s*"(.*?)(?<!\\)"', raw, re.DOTALL)
            template_text = tmpl_match.group(1).replace('\\"', '"') if tmpl_match else ""
            logger.warning(
                "parse_template_json (%s): truncated JSON — salvaged %d fields",
                source, len(salvaged),
            )
            return {"fields": salvaged, "template_text": template_text}

    logger.warning("parse_template_json (%s): could not parse JSON: %s", source, raw[:200])
    return {"fields": [], "template_text": ""}


# ── Abstract base ──────────────────────────────────────────────────────────────

class BaseAIService(ABC):
    """Abstract base class for AI services used in dynamic field extraction."""

    @abstractmethod
    async def extract_template_fields(self, file_bytes: bytes, mime_type: str, filename: str = "") -> Dict:
        """
        Analyze a document to find all blank fields.
        Returns: { fields: [{id, label}], template_text: str }
        """
        pass

    @abstractmethod
    async def generate_template(self, description: str) -> Dict:
        """
        Generate a document template from a natural-language description.
        Returns: { fields: [{id, label}], template_text: str }
        """
        pass
