from google import genai
import os
import time
from typing import Dict
from docx import Document
from base_translator import BaseTranslator

class GeminiTranslator(BaseTranslator):
    def __init__(self):
        """Initialize Gemini client with API key from environment"""
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("GEMINI_API_KEY environment variable not set")
        self.client = genai.Client(api_key=api_key)
        self.wingdings_map = {
            '\uF0A7': '☑', '\uF0A8': '☐', 
            '\uF071': '✓', '\uF072': '✗',
            '\uF06F': '☐', '\uF0FE': '■',
        }
    
    def translate_html(self, html_content: str, target_language: str) -> Dict:
        """
        Translate HTML content while preserving formatting
        
        Args:
            html_content: The HTML string to translate
            target_language: Target language code (e.g., 'hi', 'gu', 'ml')
        
        Returns:
            Dict with translated_html and language
        """
        # Language code to name mapping
        language_names = {
            'hi': 'Hindi', 'te': 'Telugu', 'mr': 'Marathi',
            'bn': 'Bengali', 'kn': 'Kannada', 'ta': 'Tamil',
            'gu': 'Gujarati', 'or': 'Odia', 'ml': 'Malayalam',
            'pa': 'Punjabi', 'as': 'Assamese', 'rajasthani': 'Hindi'
        }
        
        target_lang_name = language_names.get(target_language, target_language)
        
        prompt = f"""You are an expert document translator specializing in corporate and regulatory documents. Translate the following HTML document to {target_lang_name}.

═══════════════════════════════════════════════════════════
SECTION A: HTML STRUCTURE PRESERVATION (ABSOLUTE RULES)
═══════════════════════════════════════════════════════════

1. **PRESERVE ALL HTML TAGS EXACTLY** - Every <div>, <p>, <span>, <table>, <tr>, <td>, <th>, <h1>-<h6>, <ul>, <ol>, <li>, <br>, <hr> must remain unchanged.
2. **PRESERVE ALL <style> BLOCKS** - Do NOT translate or modify anything inside <style>...</style> tags.
3. **PRESERVE ALL INLINE STYLES** - Every style="..." attribute must remain exactly as-is. This includes `text-align: center`, `position: absolute`, etc.
4. **PRESERVE ALL CLASS & ID ATTRIBUTES** - class="page", class="docx-header", class="docx-footer", class="docx-wrapper", etc. must NOT change.
5. **PRESERVE ALL TABLE STRUCTURE** - Maintain the EXACT number of <tr> and <td> tags. NEVER merge cells or skip rows.
6. **PRESERVE ALL IMAGES** - <img> tags must be kept exactly. DO NOT modify anything inside the `src="..."` attribute.
7. **PRESERVE [[IMG_PLACEHOLDER_N]] TOKENS** - These appear inside `src` attributes. Keep them exactly as-is (e.g., [[IMG_PLACEHOLDER_0]]).
8. **DO NOT ADD OR REMOVE ANY TAGS** - The output must have the identical tag structure as the input.

═══════════════════════════════════════════════════════════
SECTION B: WHAT TO TRANSLATE
═══════════════════════════════════════════════════════════

Translate ONLY the visible text content that appears BETWEEN HTML tags:
- Paragraph text, Headings, Table cell text, List items.
- Text content inside `docx-header` and `docx-footer` divs.

═══════════════════════════════════════════════════════════
SECTION C: DO NOT TRANSLATE
═══════════════════════════════════════════════════════════

- **Company Names:** L&T, L&T Finance, L&T Finance Limited, M/s, Pvt Ltd, Co., Ltd., Inc., Corp.
- **Financial Terms:** RBI, SEBI, GST, PAN, TDS, EMI, etc.
- **Technical IDs:** Any string inside brackets like [[...]]

Return ONLY the translated HTML. No explanations, no markdown code blocks.

HTML to translate:
{html_content}

Translated HTML:"""
        
        for attempt in range(4):
            try:
                response = self.client.models.generate_content(
                    model="gemini-2.0-flash",
                    contents=prompt
                )
                break
            except Exception as e:
                if '429' in str(e) or 'RESOURCE_EXHAUSTED' in str(e):
                    wait = 30 * (attempt + 1)  # 30s, 60s, 90s
                    print(f"[Gemini] Rate limit hit, retrying in {wait}s (attempt {attempt+1}/4)")
                    time.sleep(wait)
                else:
                    raise
        else:
            raise RuntimeError("Gemini rate limit: all retries exhausted")

        translated_html = response.text.strip()

        # Remove markdown code blocks if Gemini wrapped the response
        if translated_html.startswith("```"):
            lines = translated_html.split('\n')
            if len(lines) > 2:
                translated_html = '\n'.join(lines[1:-1])

        return {
            "translated_html": translated_html,
            "language": target_language
        }

    def translate_docx(self, docx_path: str, target_language: str, output_path: str) -> str:
        """Translate DOCX natively using python-docx with full layout preservation"""
        from docx import Document
        doc = Document(docx_path)
        
        language_names = {
            'hi': 'Hindi', 'te': 'Telugu', 'mr': 'Marathi', 'kn': 'Kannada',
            'ta': 'Tamil', 'bn': 'Bengali', 'gu': 'Gujarati', 'or': 'Odia',
            'pa': 'Punjabi', 'as': 'Assamese', 'ml': 'Malayalam', 'rajasthani': 'Hindi'
        }
        target_lang_name = language_names.get(target_language.lower(), target_language)

        # 1. Collect all unique translatable text blocks
        text_blocks = set()

        def collect_from_container(container):
            for para in container.paragraphs:
                if para.text.strip():
                    text = para.text
                    # Map Wingdings characters before sending to LLM (Fix #1 from V3 report)
                    for char, mapped in self.wingdings_map.items():
                        text = text.replace(char, mapped)
                    text_blocks.add(text)
            
            # Recurse into tables which may contain more paragraphs
            tables = getattr(container, 'tables', [])
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        collect_from_container(cell)

        # Body content
        collect_from_container(doc)
        
        # Headers/Footers content (all sections and all types)
        for section in doc.sections:
            # Check all header/footer types
            header_footer_objects = [
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer
            ]
            for obj in header_footer_objects:
                if obj: collect_from_container(obj)

        blocks = list(text_blocks)
        if not blocks:
            doc.save(output_path)
            return output_path

        # 2. Batch translation (parallel fan-out)
        batch_size = 25
        batches = [blocks[i:i+batch_size] for i in range(0, len(blocks), batch_size)]

        def _translate_batch(batch):
            try:
                prompt = f"""You are a professional document translator specializing in corporate legal and regulatory documents.
Translate the following Word document segments to {target_lang_name}.

CRITICAL RULES:
1. Translate ONLY the text content provided.
2. Return precisely the same number of segments as input.
3. Separate each translated segment with a single newline (\n).
4. DO NOT add any explanations, numeric indices, or formatting.
5. PRESERVE company names exactly: "L&T", "L&T Finance", "M/s", "Pvt Ltd", "Ltd.", "Co.", "Limited".
6. PRESERVE numbers, dates, and technical abbreviations (RBI, SEBI, GST, PAN).
7. PRESERVE and INCLUDE checkboxes (☑, ☐) and symbols (✓, ✗) exactly as they appear in the source segments.

Segments to translate:
---
{chr(10).join(batch)}
---

Translated segments:"""
                for attempt in range(4):
                    try:
                        response = self.client.models.generate_content(
                            model="gemini-2.0-flash",
                            contents=prompt
                        )
                        break
                    except Exception as e:
                        if '429' in str(e) or 'RESOURCE_EXHAUSTED' in str(e):
                            wait = 30 * (attempt + 1)
                            print(f"[Gemini] Rate limit hit in batch, retrying in {wait}s (attempt {attempt+1}/4)")
                            time.sleep(wait)
                        else:
                            raise
                else:
                    raise RuntimeError("Gemini rate limit: all retries exhausted")
                raw_output = response.text.strip()
                if raw_output.startswith("```"):
                    lines = raw_output.split('\n')
                    if len(lines) > 2:
                        raw_output = '\n'.join(lines[1:-1]).strip()
                batch_translated = [t.strip() for t in raw_output.split('\n') if t.strip()]
                return dict(zip(batch, batch_translated))
            except Exception as e:
                print(f"[Gemini] Batch translation failed, keeping originals: {e}")
                return {seg: seg for seg in batch}  # fallback: keep original text

        # Gemini free tier: ~15 RPM limit — pace batches with 5s gap to avoid rate limits
        translated_map = {}
        for i, batch in enumerate(batches):
            translated_map.update(_translate_batch(batch))
            if i < len(batches) - 1:   # no sleep after the last batch
                time.sleep(5)

        from docx.oxml.ns import qn

        def apply_to_container(container):
            is_hindi = target_language.lower() in ['hi', 'rajasthani']
            
            for para in container.paragraphs:
                # Need to map the para.text for key matching since we mapped text_blocks
                lookup_text = para.text
                for char, mapped in self.wingdings_map.items():
                    lookup_text = lookup_text.replace(char, mapped)
                
                if lookup_text in translated_map:
                    translated_text = translated_map[lookup_text]
                    if para.runs:
                        # PROBST preservation: Identify runs containing drawings or picts (images)
                        text_run_indices = []
                        for i, run in enumerate(para.runs):
                            has_drawing = run._element.findall('.//' + qn('w:drawing'))
                            has_pict = run._element.findall('.//' + qn('w:pict'))
                            if not has_drawing and not has_pict:
                                text_run_indices.append(i)
                        
                        if text_run_indices:
                            # Update the first text-only run and clear others
                            main_run = para.runs[text_run_indices[0]]
                            main_run.text = translated_text
                            
                            # Enforce Hindi font if applicable
                            if is_hindi:
                                main_run.font.name = 'Anek Devanagari'
                                # Explicitly set complex script font in XML
                                rPr = main_run._element.get_or_add_rPr()
                                rFonts = rPr.get_or_add_rFonts()
                                rFonts.set(qn('w:cs'), 'Anek Devanagari')
                                rFonts.set(qn('w:hAnsi'), 'Anek Devanagari')
                                rFonts.set(qn('w:ascii'), 'Anek Devanagari')

                            for idx in text_run_indices[1:]:
                                para.runs[idx].text = ""
                        else:
                            # Add a new run if all existing ones have drawings/picts
                            new_run = para.add_run(translated_text)
                            if is_hindi:
                                new_run.font.name = 'Anek Devanagari'
                                rPr = new_run._element.get_or_add_rPr()
                                rFonts = rPr.get_or_add_rFonts()
                                rFonts.set(qn('w:cs'), 'Anek Devanagari')
                    else:
                        new_run = para.add_run(translated_text)
                        if is_hindi:
                            new_run.font.name = 'Anek Devanagari'
                            rPr = new_run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:cs'), 'Anek Devanagari')
            
            tables = getattr(container, 'tables', [])
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        apply_to_container(cell)

        # Apply to all layers
        apply_to_container(doc)
        for section in doc.sections:
            for obj in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
                if obj: apply_to_container(obj)

        doc.save(output_path)
        return output_path
