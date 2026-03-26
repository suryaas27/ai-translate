import anthropic
import os
import re
import time
from typing import Dict, Tuple
from docx import Document
from base_translator import BaseTranslator


class AnthropicTranslator(BaseTranslator):
    # Terms that must NEVER be translated — ordered longest-first to avoid partial matches
    PROTECTED_TERMS = [
        "L&T Finance Holdings Limited", "L&T Finance Holdings",
        "L&T Finance Limited", "L&T Housing Finance Limited",
        "L&T Housing Finance", "L&T Finance", "L&T",
        # Entity suffixes
        "Pvt. Ltd.", "Pvt Ltd", "Ltd.", "Co.", "Inc.", "Corp.", "Limited", "M/s",
        # Regulatory / financial abbreviations
        "NACH", "CIBIL", "NEFT", "RTGS", "MSME", "NBFC", "FOIR",
        "RBI", "SEBI", "GST", "PAN", "TDS", "EMI", "UPI", "KYC",
        "NPA", "MOU", "LOA", "NOC", "CIN", "DIN", "LLPIN", "SRN",
        "ROI", "IRR", "APR",
    ]

    def __init__(self):
        """Initialize Anthropic client with API key from environment"""
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY environment variable not set")
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = os.getenv("ANTHROPIC_MODEL", "claude-haiku-4-5-20251001")
        self.max_retries = 3
        self.wingdings_map = {
            '\uF0A7': '☑', '\uF0A8': '☐',
            '\uF071': '✓', '\uF072': '✗',
            '\uF06F': '☐', '\uF0FE': '■',
        }

    def _protect_terms(self, text: str) -> Tuple[str, Dict[str, str]]:
        """Replace protected terms with __TERM_N__ tokens so the LLM cannot alter them."""
        term_map: Dict[str, str] = {}
        for i, term in enumerate(self.PROTECTED_TERMS):
            if term in text:
                token = f"__TERM_{i}__"
                term_map[token] = term
                text = text.replace(term, token)
        # Also protect [[...]] placeholder tokens
        bracket_tokens = re.findall(r'\[\[.*?\]\]', text)
        for j, bt in enumerate(bracket_tokens):
            token = f"__BRACKET_{j}__"
            if token not in term_map:
                term_map[token] = bt
                text = text.replace(bt, token, 1)
        return text, term_map

    def _restore_terms(self, text: str, term_map: Dict[str, str]) -> str:
        """Restore __TERM_N__ / __BRACKET_N__ tokens back to their original values."""
        for token, original in term_map.items():
            text = text.replace(token, original)
        return text

    def _call_with_retry(self, **kwargs):
        """Call the Anthropic API with exponential backoff on 529 overloaded errors."""
        for attempt in range(self.max_retries):
            try:
                return self.client.messages.create(**kwargs)
            except anthropic.APIStatusError as e:
                if e.status_code == 529 and attempt < self.max_retries - 1:
                    wait = 2 ** attempt  # 1s, 2s, 4s
                    print(f"[Anthropic] Overloaded (529), retrying in {wait}s (attempt {attempt + 1}/{self.max_retries})")
                    time.sleep(wait)
                else:
                    raise

    def translate_html(self, html_content: str, target_language: str) -> Dict:
        """
        Translate HTML content while preserving formatting using Claude

        Args:
            html_content: The HTML string to translate
            target_language: Target language code (e.g., 'hi', 'te', 'mr')

        Returns:
            Dict with translated_html and language
        """
        language_names = {
            'hi': 'Hindi', 'te': 'Telugu', 'mr': 'Marathi',
            'bn': 'Bengali', 'kn': 'Kannada', 'ta': 'Tamil',
            'gu': 'Gujarati', 'or': 'Odia', 'ml': 'Malayalam',
            'pa': 'Punjabi', 'as': 'Assamese', 'rajasthani': 'Hindi'
        }

        target_lang_name = language_names.get(target_language, target_language)

        # Protect known terms before sending to LLM
        protected_html, term_map = self._protect_terms(html_content)

        token_note = ""
        if term_map:
            token_note = (
                "\n⚠️  TOKEN PRESERVATION: The input contains __TERM_N__ and __BRACKET_N__ tokens "
                "(e.g., __TERM_0__, __TERM_1__, __BRACKET_0__). These are placeholders for proper "
                "nouns and IDs. You MUST copy them into the output EXACTLY as-is — never translate, "
                "modify, or remove them.\n"
            )

        prompt = f"""You are an expert document translator specializing in corporate and regulatory documents. Translate the following HTML document to {target_lang_name}.
{token_note}
═══════════════════════════════════════════════════════════
SECTION A: HTML STRUCTURE PRESERVATION (ABSOLUTE RULES)
═══════════════════════════════════════════════════════════

1. **PRESERVE ALL HTML TAGS EXACTLY** - Every <div>, <p>, <span>, <table>, <tr>, <td>, <th>, <h1>-<h6>, <ul>, <ol>, <li>, <br>, <hr> must remain unchanged.
2. **PRESERVE ALL <style> BLOCKS** - Do NOT translate or modify anything inside <style>...</style> tags.
3. **PRESERVE ALL INLINE STYLES** - Every style="..." attribute must remain exactly as-is. This includes `text-align: center`, `position: absolute`, etc.
4. **PRESERVE ALL CLASS & ID ATTRIBUTES** - class="page", class="docx-header", class="docx-footer", class="docx-wrapper", etc. must NOT change.
5. **PRESERVE ALL TABLE STRUCTURE** - Maintain the EXACT number of <tr> and <td> tags. NEVER merge cells or skip rows.
6. **PRESERVE ALL IMAGES** - <img> tags must be kept exactly. DO NOT modify anything inside the `src="..."` attribute.
7. **PRESERVE [[IMG_PLACEHOLDER_N]] TOKENS** - These appear inside `src` attributes. Keep them exactly as-is (e.g., [[IMG_PLACEHOLDER_0]]).
8. **DO NOT ADD OR REMOVE ANY TAGS** - The output must have the identical tag structure as the input.

═══════════════════════════════════════════════════════════
SECTION B: WHAT TO TRANSLATE
═══════════════════════════════════════════════════════════

Translate ONLY the visible text content that appears BETWEEN HTML tags:
- Paragraph text, Headings, Table cell text, List items.
- Text content inside `docx-header` and `docx-footer` divs.

═══════════════════════════════════════════════════════════
SECTION C: ABSOLUTE DO-NOT-TRANSLATE LIST
═══════════════════════════════════════════════════════════

UNDER NO CIRCUMSTANCES translate ANY of the following. They must appear in the output EXACTLY as they appear in the input:

**Company / Entity Names (keep in Roman script, word-for-word):**
L&T, L&T Finance, L&T Finance Limited, L&T Finance Holdings, L&T Finance Holdings Limited,
L&T Housing Finance, L&T Housing Finance Limited, M/s, Pvt Ltd, Pvt. Ltd., Ltd., Co., Inc., Corp., Limited

**Regulatory & Financial Abbreviations (keep as uppercase Roman letters):**
RBI, SEBI, GST, PAN, TDS, EMI, NACH, CIBIL, KYC, NEFT, RTGS, UPI, MSME, NPA, NBFC,
MOU, LOA, NOC, CIN, DIN, LLPIN, SRN, ROI, IRR, APR, FOIR

**Placeholder Tokens (copy byte-for-byte):**
Any token matching __TERM_N__ or __BRACKET_N__ (e.g., __TERM_0__, __BRACKET_2__)
Any string matching [[...]] (e.g., [[IMG_PLACEHOLDER_0]])

❌ WRONG:  "एल एंड टी फाइनेंस लिमिटेड"  (translated company name)
✅ CORRECT: "L&T Finance Limited"            (kept in Roman script)

❌ WRONG:  "आरबीआई"  (translated abbreviation)
✅ CORRECT: "RBI"     (kept as-is)

Return ONLY the translated HTML. No explanations, no markdown code blocks.

HTML to translate:
{protected_html}

Translated HTML:"""

        response = self._call_with_retry(
            model=self.model,
            max_tokens=8096,
            system=(
                "You are an expert corporate document translator. Your ONLY job is to translate "
                "text content between HTML tags while preserving ALL HTML structure, CSS, styles, "
                "images, tables, and layout EXACTLY as given. You specialize in Indian regulatory "
                "and financial documents. CRITICAL: Company names (L&T, L&T Finance, etc.), "
                "regulatory abbreviations (RBI, SEBI, GST, PAN, TDS, EMI, NACH, CIBIL, KYC, NEFT, "
                "RTGS, UPI, MSME, NBFC), and placeholder tokens (__TERM_N__, __BRACKET_N__, [[...]]) "
                "must NEVER be translated — copy them verbatim into the output."
            ),
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )

        translated_html = response.content[0].text.strip()

        # Remove markdown code blocks if Claude wrapped the response
        if translated_html.startswith("```"):
            lines = translated_html.split('\n')
            if len(lines) > 2:
                start_idx = 1 if lines[0].startswith("```") else 0
                end_idx = -1 if lines[-1].strip() == "```" else len(lines)
                translated_html = '\n'.join(lines[start_idx:end_idx]).strip()

        # Restore protected terms
        translated_html = self._restore_terms(translated_html, term_map)

        return {
            "translated_html": translated_html,
            "language": target_language
        }

    def translate_docx(self, docx_path: str, target_language: str, output_path: str) -> str:
        """Translate DOCX natively using python-docx with full layout preservation"""
        doc = Document(docx_path)

        language_names = {
            'hi': 'Hindi', 'te': 'Telugu', 'mr': 'Marathi', 'kn': 'Kannada',
            'ta': 'Tamil', 'bn': 'Bengali', 'gu': 'Gujarati', 'or': 'Odia',
            'pa': 'Punjabi', 'as': 'Assamese', 'ml': 'Malayalam', 'rajasthani': 'Hindi'
        }
        target_lang_name = language_names.get(target_language.lower(), target_language)

        # 1. Collect all unique translatable text blocks
        text_blocks = set()

        def collect_from_container(container):
            for para in container.paragraphs:
                if para.text.strip():
                    text = para.text
                    for char, mapped in self.wingdings_map.items():
                        text = text.replace(char, mapped)
                    text_blocks.add(text)
            tables = getattr(container, 'tables', [])
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        collect_from_container(cell)

        collect_from_container(doc)
        for section in doc.sections:
            for obj in [
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer
            ]:
                if obj:
                    collect_from_container(obj)

        blocks = list(text_blocks)
        if not blocks:
            doc.save(output_path)
            return output_path

        # 2. Batch translation (parallel fan-out with per-batch fallback)
        batch_size = 20
        batches = [blocks[i:i + batch_size] for i in range(0, len(blocks), batch_size)]

        def _translate_batch(batch):
            try:
                # Protect terms in each segment
                protected_batch = []
                batch_term_maps = []
                for seg in batch:
                    p_seg, t_map = self._protect_terms(seg)
                    protected_batch.append(p_seg)
                    batch_term_maps.append(t_map)

                prompt = f"""Translate each text segment to {target_lang_name}.

Return EXACTLY {len(batch)} lines — one translation per input segment.
- __TERM_N__ and __BRACKET_N__ tokens → copy verbatim, never translate.
- Checkboxes (☑ ☐ ✓ ✗ ■) → copy exactly as-is.
- No explanations, numbering, or markdown.

Segments:
---
{chr(10).join(protected_batch)}
---

Translated ({len(batch)} lines):"""
                response = self._call_with_retry(
                    model=self.model,
                    max_tokens=4096,
                    system=(
                        f"Professional corporate translator. Output exactly {len(batch)} lines "
                        f"matching the input count. Never add commentary or formatting. "
                        f"__TERM_N__ and __BRACKET_N__ tokens must be copied verbatim."
                    ),
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0
                )
                raw_output = response.content[0].text.strip()
                if raw_output.startswith("```"):
                    lines = raw_output.split('\n')
                    if len(lines) > 2:
                        raw_output = '\n'.join(lines[1:-1]).strip()
                batch_translated_raw = [t.strip() for t in raw_output.split('\n') if t.strip()]

                # Restore protected terms in each translated segment
                batch_translated = []
                for idx, translated_seg in enumerate(batch_translated_raw):
                    if idx < len(batch_term_maps):
                        translated_seg = self._restore_terms(translated_seg, batch_term_maps[idx])
                    batch_translated.append(translated_seg)

                return dict(zip(batch, batch_translated))
            except Exception as e:
                print(f"[Anthropic] Batch translation failed, keeping originals: {e}")
                return {seg: seg for seg in batch}  # fallback: keep original text

        # Anthropic has strict TPM limits — run batches sequentially to avoid rate limit collisions
        translated_map = {}
        for batch in batches:
            translated_map.update(_translate_batch(batch))

        from docx.oxml.ns import qn

        def apply_to_container(container):
            is_hindi = target_language.lower() in ['hi', 'rajasthani']

            for para in container.paragraphs:
                lookup_text = para.text
                for char, mapped in self.wingdings_map.items():
                    lookup_text = lookup_text.replace(char, mapped)

                if lookup_text in translated_map:
                    translated_text = translated_map[lookup_text]
                    if para.runs:
                        text_run_indices = []
                        for idx, run in enumerate(para.runs):
                            has_drawing = run._element.findall('.//' + qn('w:drawing'))
                            has_pict = run._element.findall('.//' + qn('w:pict'))
                            if not has_drawing and not has_pict:
                                text_run_indices.append(idx)

                        if text_run_indices:
                            main_run = para.runs[text_run_indices[0]]
                            main_run.text = translated_text

                            if is_hindi:
                                main_run.font.name = 'Anek Devanagari'
                                rPr = main_run._element.get_or_add_rPr()
                                rFonts = rPr.get_or_add_rFonts()
                                rFonts.set(qn('w:cs'), 'Anek Devanagari')
                                rFonts.set(qn('w:hAnsi'), 'Anek Devanagari')
                                rFonts.set(qn('w:ascii'), 'Anek Devanagari')

                            for idx in text_run_indices[1:]:
                                para.runs[idx].text = ""
                        else:
                            new_run = para.add_run(translated_text)
                            if is_hindi:
                                new_run.font.name = 'Anek Devanagari'
                                rPr = new_run._element.get_or_add_rPr()
                                rFonts = rPr.get_or_add_rFonts()
                                rFonts.set(qn('w:cs'), 'Anek Devanagari')
                    else:
                        new_run = para.add_run(translated_text)
                        if is_hindi:
                            new_run.font.name = 'Anek Devanagari'
                            rPr = new_run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:cs'), 'Anek Devanagari')

            tables = getattr(container, 'tables', [])
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        apply_to_container(cell)

        apply_to_container(doc)
        for section in doc.sections:
            for obj in [
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer
            ]:
                if obj:
                    apply_to_container(obj)

        doc.save(output_path)
        return output_path
