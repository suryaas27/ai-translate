import boto3
import botocore.config
import botocore.exceptions
import os
import re
import time
from typing import Dict, Tuple
from docx import Document
from base_translator import BaseTranslator


class BedrockTranslator(BaseTranslator):
    """Translates via AWS Bedrock (Claude). Used as the 'anthropic' provider in server flow."""

    PROTECTED_TERMS = [
        "L&T Finance Holdings Limited", "L&T Finance Holdings",
        "L&T Finance Limited", "L&T Housing Finance Limited",
        "L&T Housing Finance", "L&T Finance", "L&T",
        "Pvt. Ltd.", "Pvt Ltd", "Ltd.", "Co.", "Inc.", "Corp.", "Limited", "M/s",
        "NACH", "CIBIL", "NEFT", "RTGS", "MSME", "NBFC", "FOIR",
        "RBI", "SEBI", "GST", "PAN", "TDS", "EMI", "UPI", "KYC",
        "NPA", "MOU", "LOA", "NOC", "CIN", "DIN", "LLPIN", "SRN",
        "ROI", "IRR", "APR",
    ]

    LANGUAGE_NAMES = {
        'hi': 'Hindi', 'te': 'Telugu', 'mr': 'Marathi', 'kn': 'Kannada',
        'ta': 'Tamil', 'bn': 'Bengali', 'gu': 'Gujarati', 'or': 'Odia',
        'pa': 'Punjabi', 'as': 'Assamese', 'ml': 'Malayalam', 'rajasthani': 'Hindi',
    }

    def __init__(self):
        region = os.getenv("AWS_DEFAULT_REGION", "ap-south-1")
        self.model_id = os.getenv("BEDROCK_MODEL", "global.anthropic.claude-haiku-4-5-20251001-v1:0")
        self.max_tokens = 8096
        self.max_retries = 3
        self.client = boto3.client(
            "bedrock-runtime",
            region_name=region,
            config=botocore.config.Config(
                read_timeout=300,
                connect_timeout=10,
                retries={"max_attempts": 0},
            ),
        )
        self.wingdings_map = {
            '\uF0A7': '☑', '\uF0A8': '☐',
            '\uF071': '✓', '\uF072': '✗',
            '\uF06F': '☐', '\uF0FE': '■',
        }
        print(f"[BedrockTranslator] Initialized: model={self.model_id}, region={region}")

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _protect_terms(self, text: str) -> Tuple[str, Dict[str, str]]:
        term_map: Dict[str, str] = {}
        for i, term in enumerate(self.PROTECTED_TERMS):
            if term in text:
                token = f"__TERM_{i}__"
                term_map[token] = term
                text = text.replace(term, token)
        bracket_tokens = re.findall(r'\[\[.*?\]\]', text)
        for j, bt in enumerate(bracket_tokens):
            token = f"__BRACKET_{j}__"
            if token not in term_map:
                term_map[token] = bt
                text = text.replace(bt, token, 1)
        return text, term_map

    def _restore_terms(self, text: str, term_map: Dict[str, str]) -> str:
        for token, original in term_map.items():
            text = text.replace(token, original)
        return text

    def _call_bedrock(self, system_prompt: str, user_message: str) -> Tuple[str, int, int]:
        """Returns (output_text, input_tokens, output_tokens)."""
        for attempt in range(self.max_retries):
            try:
                response = self.client.converse(
                    modelId=self.model_id,
                    system=[{"text": system_prompt}],
                    messages=[{"role": "user", "content": [{"text": user_message}]}],
                    inferenceConfig={"maxTokens": self.max_tokens, "temperature": 0.0},
                )
                output = response["output"]["message"]["content"][0]["text"]
                if response.get("stopReason") == "max_tokens":
                    print(f"[Bedrock] Warning: hit max_tokens ({self.max_tokens}), output may be truncated")
                usage = response.get("usage", {})
                input_tokens = usage.get("inputTokens", 0)
                output_tokens = usage.get("outputTokens", 0)
                return output, input_tokens, output_tokens
            except botocore.exceptions.ClientError as e:
                code = e.response["Error"]["Code"]
                if code in ("ThrottlingException", "ServiceUnavailableException") and attempt < self.max_retries - 1:
                    wait = 2 ** attempt
                    print(f"[Bedrock] {code}, retrying in {wait}s (attempt {attempt+1}/{self.max_retries})")
                    time.sleep(wait)
                else:
                    raise
        raise RuntimeError(f"[Bedrock] API failed after {self.max_retries} retries")

    # ------------------------------------------------------------------
    # BaseTranslator interface
    # ------------------------------------------------------------------

    def translate_html(self, html_content: str, target_language: str) -> Dict:
        target_lang_name = self.LANGUAGE_NAMES.get(target_language, target_language)
        protected_html, term_map = self._protect_terms(html_content)

        token_note = ""
        if term_map:
            token_note = (
                "\n⚠️  TOKEN PRESERVATION: The input contains __TERM_N__ and __BRACKET_N__ tokens "
                "(e.g., __TERM_0__, __TERM_1__, __BRACKET_0__). These are placeholders for proper "
                "nouns and IDs. You MUST copy them into the output EXACTLY as-is — never translate, "
                "modify, or remove them.\n"
            )

        prompt = f"""You are an expert document translator specializing in corporate and regulatory documents. Translate the following HTML document to {target_lang_name}.
{token_note}
═══════════════════════════════════════════════════════════
SECTION A: HTML STRUCTURE PRESERVATION (ABSOLUTE RULES)
═══════════════════════════════════════════════════════════

1. **PRESERVE ALL HTML TAGS EXACTLY** - Every <div>, <p>, <span>, <table>, <tr>, <td>, <th>, <h1>-<h6>, <ul>, <ol>, <li>, <br>, <hr> must remain unchanged.
2. **PRESERVE ALL <style> BLOCKS** - Do NOT translate or modify anything inside <style>...</style> tags.
3. **PRESERVE ALL INLINE STYLES** - Every style="..." attribute must remain exactly as-is.
4. **PRESERVE ALL CLASS & ID ATTRIBUTES** - class="page", class="docx-header", class="docx-footer", etc. must NOT change.
5. **PRESERVE ALL TABLE STRUCTURE** - Maintain the EXACT number of <tr> and <td> tags. NEVER merge cells or skip rows.
6. **PRESERVE ALL IMAGES** - <img> tags must be kept exactly. DO NOT modify the `src="..."` attribute.
7. **PRESERVE [[IMG_PLACEHOLDER_N]] TOKENS** - Keep them exactly as-is (e.g., [[IMG_PLACEHOLDER_0]]).
8. **DO NOT ADD OR REMOVE ANY TAGS** - The output must have the identical tag structure as the input.

═══════════════════════════════════════════════════════════
SECTION B: WHAT TO TRANSLATE
═══════════════════════════════════════════════════════════

Translate ONLY the visible text content that appears BETWEEN HTML tags:
- Paragraph text, Headings, Table cell text, List items.
- Text content inside `docx-header` and `docx-footer` divs.

═══════════════════════════════════════════════════════════
SECTION C: ABSOLUTE DO-NOT-TRANSLATE LIST
═══════════════════════════════════════════════════════════

UNDER NO CIRCUMSTANCES translate ANY of the following:

**Company / Entity Names (keep in Roman script):**
L&T, L&T Finance, L&T Finance Limited, L&T Finance Holdings, L&T Finance Holdings Limited,
L&T Housing Finance, L&T Housing Finance Limited, M/s, Pvt Ltd, Pvt. Ltd., Ltd., Co., Inc., Corp., Limited

**Regulatory & Financial Abbreviations:**
RBI, SEBI, GST, PAN, TDS, EMI, NACH, CIBIL, KYC, NEFT, RTGS, UPI, MSME, NPA, NBFC,
MOU, LOA, NOC, CIN, DIN, LLPIN, SRN, ROI, IRR, APR, FOIR

**Placeholder Tokens (copy byte-for-byte):**
Any token matching __TERM_N__ or __BRACKET_N__ (e.g., __TERM_0__, __BRACKET_2__)
Any string matching [[...]] (e.g., [[IMG_PLACEHOLDER_0]])

⚠️  CRITICAL TOKEN RULE: NEVER invent or create new __TERM_N__ or __BRACKET_N__ tokens.
Only copy tokens that ALREADY EXIST verbatim in the input HTML.
Words like "Company", "Borrower", "Lender", etc. must be TRANSLATED normally — do NOT replace them with tokens.

❌ WRONG:  Replacing "Company" with "__TERM_1__"   →  ✅ CORRECT: Translate "Company" to the target language
❌ WRONG:  "एल एंड टी फाइनेंस लिमिटेड"           →  ✅ CORRECT: "L&T Finance Limited"
❌ WRONG:  "आरबीआई"                                →  ✅ CORRECT: "RBI"

Return ONLY the translated HTML. No explanations, no markdown code blocks.

HTML to translate:
{protected_html}

Translated HTML:"""

        system_prompt = (
            "You are an expert corporate document translator. Your ONLY job is to translate "
            "text content between HTML tags while preserving ALL HTML structure, CSS, styles, "
            "images, tables, and layout EXACTLY as given. You specialize in Indian regulatory "
            "and financial documents. CRITICAL: Company names (L&T, L&T Finance, etc.), "
            "regulatory abbreviations (RBI, SEBI, GST, PAN, TDS, EMI, NACH, CIBIL, KYC, NEFT, "
            "RTGS, UPI, MSME, NBFC), and placeholder tokens (__TERM_N__, __BRACKET_N__, [[...]]) "
            "must NEVER be translated — copy them verbatim into the output. "
            "NEVER invent new __TERM_N__ tokens — only copy tokens that already appear in the input. "
            "Generic legal words like 'Company', 'Borrower', 'Lender' must be translated normally."
        )

        translated_html, input_tokens, output_tokens = self._call_bedrock(system_prompt, prompt)

        if translated_html.startswith("```"):
            lines = translated_html.split('\n')
            if len(lines) > 2:
                start_idx = 1 if lines[0].startswith("```") else 0
                end_idx = -1 if lines[-1].strip() == "```" else len(lines)
                translated_html = '\n'.join(lines[start_idx:end_idx]).strip()

        translated_html = self._restore_terms(translated_html, term_map)

        print(f"[Bedrock] Token usage — input: {input_tokens:,}, output: {output_tokens:,}, total: {input_tokens + output_tokens:,}")

        return {"translated_html": translated_html, "language": target_language}

    def translate_docx(self, docx_path: str, target_language: str, output_path: str) -> str:
        doc = Document(docx_path)
        target_lang_name = self.LANGUAGE_NAMES.get(target_language.lower(), target_language)

        text_blocks = set()

        def collect_from_container(container):
            for para in container.paragraphs:
                if para.text.strip():
                    text = para.text
                    for char, mapped in self.wingdings_map.items():
                        text = text.replace(char, mapped)
                    text_blocks.add(text)
            for table in getattr(container, 'tables', []):
                for row in table.rows:
                    for cell in row.cells:
                        collect_from_container(cell)

        collect_from_container(doc)
        for section in doc.sections:
            for obj in [
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer,
            ]:
                if obj:
                    collect_from_container(obj)

        blocks = list(text_blocks)
        if not blocks:
            doc.save(output_path)
            return output_path

        batch_size = 20
        batches = [blocks[i:i + batch_size] for i in range(0, len(blocks), batch_size)]

        def _translate_batch(batch):
            try:
                protected_batch = []
                batch_term_maps = []
                for seg in batch:
                    p_seg, t_map = self._protect_terms(seg)
                    protected_batch.append(p_seg)
                    batch_term_maps.append(t_map)

                prompt = f"""Translate each text segment to {target_lang_name}.

Return EXACTLY {len(batch)} lines — one translation per input segment.
- __TERM_N__ and __BRACKET_N__ tokens → copy verbatim, never translate.
- Checkboxes (☑ ☐ ✓ ✗ ■) → copy exactly as-is.
- No explanations, numbering, or markdown.

Segments:
---
{chr(10).join(protected_batch)}
---

Translated ({len(batch)} lines):"""

                system_prompt = (
                    f"Professional corporate translator. Output exactly {len(batch)} lines "
                    f"matching the input count. Never add commentary or formatting. "
                    f"__TERM_N__ and __BRACKET_N__ tokens must be copied verbatim."
                )

                raw_output, in_tok, out_tok = self._call_bedrock(system_prompt, prompt)
                if raw_output.startswith("```"):
                    lines = raw_output.split('\n')
                    if len(lines) > 2:
                        raw_output = '\n'.join(lines[1:-1]).strip()

                batch_translated_raw = [t.strip() for t in raw_output.split('\n') if t.strip()]
                batch_translated = []
                for idx, translated_seg in enumerate(batch_translated_raw):
                    if idx < len(batch_term_maps):
                        translated_seg = self._restore_terms(translated_seg, batch_term_maps[idx])
                    batch_translated.append(translated_seg)

                return dict(zip(batch, batch_translated)), in_tok, out_tok
            except Exception as e:
                print(f"[Bedrock] Batch translation failed, keeping originals: {e}")
                return {seg: seg for seg in batch}, 0, 0

        doc_name = os.path.basename(docx_path)
        total_input_tokens = 0
        total_output_tokens = 0
        translated_map = {}
        for batch in batches:
            batch_result, in_tok, out_tok = _translate_batch(batch)
            translated_map.update(batch_result)
            total_input_tokens += in_tok
            total_output_tokens += out_tok

        from docx.oxml.ns import qn

        def apply_to_container(container):
            is_hindi = target_language.lower() in ['hi', 'rajasthani']
            for para in container.paragraphs:
                lookup_text = para.text
                for char, mapped in self.wingdings_map.items():
                    lookup_text = lookup_text.replace(char, mapped)
                if lookup_text not in translated_map:
                    continue
                translated_text = translated_map[lookup_text]
                if para.runs:
                    text_run_indices = [
                        i for i, run in enumerate(para.runs)
                        if not run._element.findall('.//' + qn('w:drawing'))
                        and not run._element.findall('.//' + qn('w:pict'))
                    ]
                    if text_run_indices:
                        main_run = para.runs[text_run_indices[0]]
                        main_run.text = translated_text
                        if is_hindi:
                            main_run.font.name = 'Anek Devanagari'
                            rPr = main_run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:cs'), 'Anek Devanagari')
                            rFonts.set(qn('w:hAnsi'), 'Anek Devanagari')
                            rFonts.set(qn('w:ascii'), 'Anek Devanagari')
                        for idx in text_run_indices[1:]:
                            para.runs[idx].text = ""
                    else:
                        new_run = para.add_run(translated_text)
                        if is_hindi:
                            new_run.font.name = 'Anek Devanagari'
                            rPr = new_run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:cs'), 'Anek Devanagari')
                else:
                    new_run = para.add_run(translated_text)
                    if is_hindi:
                        new_run.font.name = 'Anek Devanagari'
                        rPr = new_run._element.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:cs'), 'Anek Devanagari')

            for table in getattr(container, 'tables', []):
                for row in table.rows:
                    for cell in row.cells:
                        apply_to_container(cell)

        apply_to_container(doc)
        for section in doc.sections:
            for obj in [
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer,
            ]:
                if obj:
                    apply_to_container(obj)

        print(
            f"[Bedrock] [{doc_name}] Token usage — "
            f"input: {total_input_tokens:,}, output: {total_output_tokens:,}, "
            f"total: {total_input_tokens + total_output_tokens:,}"
        )
        doc.save(output_path)
        return output_path

    def translate_pdf(self, pdf_path: str, target_language: str, output_path: str) -> str:
        """
        Translate PDF natively using PyMuPDF:
          1. Extract text at LINE level (per visual line, preserving layout)
          2. Batch translate plain text via Bedrock (same prompt as translate_docx)
          3. Redact original lines and insert translated text at same positions
        """
        import fitz  # pymupdf
        import os

        doc = fitz.open(pdf_path)
        target_lang_name = self.LANGUAGE_NAMES.get(target_language.lower(), target_language)

        # Look up a bundled font file for scripts that need non-Latin glyphs
        _font_map = {
            'hi': 'AnekDevanagari-Regular.ttf',
            'rajasthani': 'AnekDevanagari-Regular.ttf',
            'mr': 'AnekDevanagari-Regular.ttf',
            'te': 'NotoSansTelugu-Regular.ttf',
            'ta': 'NotoSansTamil-Regular.ttf',
            'kn': 'NotoSansKannada-Regular.ttf',
            'ml': 'NotoSansMalayalam-Regular.ttf',
            'gu': 'NotoSansGujarati-Regular.ttf',
            'bn': 'NotoSansBengali-Regular.ttf',
            'pa': 'NotoSansGurmukhi-Regular.ttf',
            'or': 'NotoSansOriya-Regular.ttf',
            'as': 'NotoSansBengali-Regular.ttf',
        }
        fonts_dir = os.path.join(os.path.dirname(__file__), "fonts")
        font_filename = _font_map.get(target_language.lower())
        font_path = None
        if font_filename:
            candidate = os.path.join(fonts_dir, font_filename)
            if os.path.exists(candidate):
                font_path = candidate
            else:
                print(f"[Bedrock-PDF] Font file not found: {candidate} — using fallback")

        # ------------------------------------------------------------------
        # 1. Collect unique line-level text blocks across all pages
        # ------------------------------------------------------------------
        # Key: stripped line text  →  Value: list of {page_no, bbox, spans, origin, fontsize, color}
        line_occurrences: Dict[str, list] = {}

        for page_no, page in enumerate(doc):
            raw = page.get_text("dict", flags=0)
            for block in raw.get("blocks", []):
                if block.get("type") != 0:
                    continue  # skip image blocks
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    # Concatenate all spans in the line → full line text
                    line_text = "".join(s["text"] for s in spans)
                    for char, mapped in self.wingdings_map.items():
                        line_text = line_text.replace(char, mapped)
                    line_text = line_text.strip()
                    if not line_text:
                        continue

                    # Bounding box that covers the whole line
                    line_bbox = fitz.Rect(line["bbox"])
                    # Use origin of first span for text insertion
                    first_span = spans[0]
                    origin = fitz.Point(first_span["origin"])
                    fontsize = first_span.get("size", 11)
                    color_int = first_span.get("color", 0)
                    # Convert integer color → (r, g, b) floats
                    color = (
                        ((color_int >> 16) & 0xFF) / 255.0,
                        ((color_int >> 8) & 0xFF) / 255.0,
                        (color_int & 0xFF) / 255.0,
                    )

                    occ = {
                        "page_no": page_no,
                        "line_bbox": line_bbox,
                        "origin": origin,
                        "fontsize": fontsize,
                        "color": color,
                    }
                    line_occurrences.setdefault(line_text, []).append(occ)

        all_lines = list(line_occurrences.keys())
        if not all_lines:
            doc.save(output_path)
            doc.close()
            return output_path

        # ------------------------------------------------------------------
        # 2. Batch translate (same prompt as translate_docx)
        # ------------------------------------------------------------------
        batch_size = 20
        batches = [all_lines[i:i + batch_size] for i in range(0, len(all_lines), batch_size)]

        def _translate_batch(batch):
            try:
                protected_batch, batch_term_maps = [], []
                for seg in batch:
                    p, tm = self._protect_terms(seg)
                    protected_batch.append(p)
                    batch_term_maps.append(tm)

                prompt = f"""Translate each text segment to {target_lang_name}.

Return EXACTLY {len(batch)} lines — one translation per input segment.
- __TERM_N__ and __BRACKET_N__ tokens → copy verbatim, never translate.
- Checkboxes (☑ ☐ ✓ ✗ ■) → copy exactly as-is.
- No explanations, numbering, or markdown.

Segments:
---
{chr(10).join(protected_batch)}
---

Translated ({len(batch)} lines):"""

                system_prompt = (
                    f"Professional corporate translator. Output exactly {len(batch)} lines "
                    f"matching the input count. Never add commentary or formatting. "
                    f"__TERM_N__ and __BRACKET_N__ tokens must be copied verbatim."
                )

                raw_output = self._call_bedrock(system_prompt, prompt)
                if raw_output.startswith("```"):
                    lines_out = raw_output.split('\n')
                    if len(lines_out) > 2:
                        raw_output = '\n'.join(lines_out[1:-1]).strip()
                batch_translated_raw = [t.strip() for t in raw_output.split('\n') if t.strip()]
                batch_translated = []
                for idx, translated_seg in enumerate(batch_translated_raw):
                    if idx < len(batch_term_maps):
                        translated_seg = self._restore_terms(translated_seg, batch_term_maps[idx])
                    batch_translated.append(translated_seg)
                return dict(zip(batch, batch_translated))
            except Exception as e:
                print(f"[Bedrock-PDF] Batch translation failed, keeping originals: {e}")
                return {seg: seg for seg in batch}

        translated_map: Dict[str, str] = {}
        for batch in batches:
            translated_map.update(_translate_batch(batch))

        print(f"[Bedrock-PDF] Translated {len(translated_map)} unique line segments")

        # ------------------------------------------------------------------
        # 3. Apply translations page by page
        # ------------------------------------------------------------------
        for page_no, page in enumerate(doc):
            raw = page.get_text("dict", flags=0)

            # Collect redactions for this page
            page_redactions = []
            for block in raw.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    line_text = "".join(s["text"] for s in spans)
                    for char, mapped in self.wingdings_map.items():
                        line_text = line_text.replace(char, mapped)
                    line_text = line_text.strip()
                    if line_text not in translated_map:
                        continue
                    translated_text = translated_map[line_text]
                    if translated_text == line_text:
                        continue  # unchanged, no need to redact

                    first_span = spans[0]
                    page_redactions.append({
                        "line_bbox": fitz.Rect(line["bbox"]),
                        "origin": fitz.Point(first_span["origin"]),
                        "fontsize": first_span.get("size", 11),
                        "color": (
                            ((first_span.get("color", 0) >> 16) & 0xFF) / 255.0,
                            ((first_span.get("color", 0) >> 8) & 0xFF) / 255.0,
                            (first_span.get("color", 0) & 0xFF) / 255.0,
                        ),
                        "translated": translated_text,
                    })

            if not page_redactions:
                continue

            # Erase original text (white fill, keep images untouched)
            for r in page_redactions:
                page.add_redact_annot(r["line_bbox"], fill=(1, 1, 1))
            page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

            # Insert translated text at same baseline origin
            for r in page_redactions:
                try:
                    kwargs = {
                        "point": r["origin"],
                        "text": r["translated"],
                        "fontsize": r["fontsize"],
                        "color": r["color"],
                    }
                    if font_path:
                        kwargs["fontfile"] = font_path
                        kwargs["fontname"] = "custom"
                    else:
                        kwargs["fontname"] = "helv"
                    page.insert_text(**kwargs)
                except Exception as e:
                    print(f"[Bedrock-PDF] insert_text failed: {e}")

        doc.save(output_path, garbage=4, deflate=True)
        doc.close()
        print(f"[Bedrock-PDF] Saved translated PDF → {output_path}")
        return output_path
