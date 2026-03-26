from openai import OpenAI
import os
import re
from typing import Dict, Tuple
from docx import Document
from base_translator import BaseTranslator

class OpenAITranslator(BaseTranslator):
    PROTECTED_TERMS = [
        "L&T Finance Holdings Limited", "L&T Finance Holdings",
        "L&T Finance Limited", "L&T Housing Finance Limited",
        "L&T Housing Finance", "L&T Finance", "L&T",
        "Pvt. Ltd.", "Pvt Ltd", "Ltd.", "Co.", "Inc.", "Corp.", "Limited", "M/s",
        "NACH", "CIBIL", "NEFT", "RTGS", "MSME", "NBFC", "FOIR",
        "RBI", "SEBI", "GST", "PAN", "TDS", "EMI", "UPI", "KYC",
        "NPA", "MOU", "LOA", "NOC", "CIN", "DIN", "LLPIN", "SRN",
        "ROI", "IRR", "APR",
    ]

    def __init__(self):
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY environment variable not set")
        self.client = OpenAI(api_key=api_key)
        self.wingdings_map = {
            '\uF0A7': '☑', '\uF0A8': '☐',
            '\uF071': '✓', '\uF072': '✗',
            '\uF06F': '☐', '\uF0FE': '■',
        }

    def _protect_terms(self, text: str) -> Tuple[str, dict]:
        term_map = {}
        for i, term in enumerate(self.PROTECTED_TERMS):
            if term in text:
                token = f"__TERM_{i}__"
                term_map[token] = term
                text = text.replace(term, token)
        for j, bt in enumerate(re.findall(r'\[\[.*?\]\]', text)):
            token = f"__BRACKET_{j}__"
            if token not in term_map:
                term_map[token] = bt
                text = text.replace(bt, token, 1)
        return text, term_map

    def _restore_terms(self, text: str, term_map: dict) -> str:
        for token, original in term_map.items():
            text = text.replace(token, original)
        return text
    
    def translate_html(self, html_content: str, target_language: str) -> Dict:
        """
        Translate HTML content while preserving formatting using OpenAI GPT-4o
        
        Args:
            html_content: The HTML string to translate
            target_language: Target language code (e.g., 'hi', 'te', 'mr')
        
        Returns:
            Dict with translated_html and language
        """
        # Language code to name mapping
        language_names = {
            'hi': 'Hindi', 'te': 'Telugu', 'mr': 'Marathi',
            'bn': 'Bengali', 'kn': 'Kannada', 'ta': 'Tamil',
            'gu': 'Gujarati', 'or': 'Odia', 'ml': 'Malayalam',
            'pa': 'Punjabi', 'as': 'Assamese', 'rajasthani': 'Hindi'
        }
        
        target_lang_name = language_names.get(target_language, target_language)
        
        prompt = f"""You are an expert document translator specializing in corporate and regulatory documents. Translate the following HTML document to {target_lang_name}.

═══════════════════════════════════════════════════════════
SECTION A: HTML STRUCTURE PRESERVATION (ABSOLUTE RULES)
═══════════════════════════════════════════════════════════

1. **PRESERVE ALL HTML TAGS EXACTLY** - Every <div>, <p>, <span>, <table>, <tr>, <td>, <th>, <h1>-<h6>, <ul>, <ol>, <li>, <br>, <hr> must remain unchanged.
2. **PRESERVE ALL <style> BLOCKS** - Do NOT translate or modify anything inside <style>...</style> tags.
3. **PRESERVE ALL INLINE STYLES** - Every style="..." attribute must remain exactly as-is. This includes `text-align: center`, `position: absolute`, etc.
4. **PRESERVE ALL CLASS & ID ATTRIBUTES** - class="page", class="docx-header", class="docx-footer", class="docx-wrapper", etc. must NOT change.
5. **PRESERVE ALL TABLE STRUCTURE** - Maintain the EXACT number of <tr> and <td> tags. NEVER merge cells or skip rows.
6. **PRESERVE ALL IMAGES** - <img> tags must be kept exactly. DO NOT modify anything inside the `src="..."` attribute.
7. **PRESERVE [[IMG_PLACEHOLDER_N]] TOKENS** - These appear inside `src` attributes. Keep them exactly as-is (e.g., [[IMG_PLACEHOLDER_0]]).
8. **DO NOT ADD OR REMOVE ANY TAGS** - The output must have the identical tag structure as the input.

═══════════════════════════════════════════════════════════
SECTION B: WHAT TO TRANSLATE
═══════════════════════════════════════════════════════════

Translate ONLY the visible text content that appears BETWEEN HTML tags:
- Paragraph text, Headings, Table cell text, List items.
- Text content inside `docx-header` and `docx-footer` divs.

═══════════════════════════════════════════════════════════
SECTION C: DO NOT TRANSLATE
═══════════════════════════════════════════════════════════

- **Company Names:** L&T, L&T Finance, L&T Finance Limited, M/s, Pvt Ltd, Co., Ltd., Inc., Corp.
- **Financial Terms:** RBI, SEBI, GST, PAN, TDS, EMI, etc.
- **Technical IDs:** Any string inside brackets like [[...]]

Return ONLY the translated HTML. No explanations, no markdown code blocks.

HTML to translate:
{html_content}

Translated HTML:"""

        response = self.client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert corporate document translator. Your ONLY job is to translate text content between HTML tags while preserving ALL HTML structure, CSS, styles, images, tables, and layout EXACTLY as given. You specialize in Indian regulatory and financial documents."},
                {"role": "user", "content": prompt}
            ],
            temperature=0
        )
        
        translated_html = response.choices[0].message.content.strip()
        
        # Remove markdown code blocks if the model wrapped the response
        if translated_html.startswith("```"):
            lines = translated_html.split('\n')
            if len(lines) > 2:
                # Find the first line that isn't the code block start
                start_idx = 1 if lines[0].startswith("```") else 0
                # Find the last line that isn't the code block end
                end_idx = -1 if lines[-1].strip() == "```" else len(lines)
                translated_html = '\n'.join(lines[start_idx:end_idx]).strip()
        
        return {
            "translated_html": translated_html,
            "language": target_language
        }

    def translate_docx(self, docx_path: str, target_language: str, output_path: str) -> str:
        """Translate DOCX natively using python-docx with full layout preservation"""
        from docx import Document
        doc = Document(docx_path)
        
        language_names = {
            'hi': 'Hindi', 'te': 'Telugu', 'mr': 'Marathi', 'kn': 'Kannada',
            'ta': 'Tamil', 'bn': 'Bengali', 'gu': 'Gujarati', 'or': 'Odia',
            'pa': 'Punjabi', 'as': 'Assamese', 'ml': 'Malayalam', 'rajasthani': 'Hindi'
        }
        target_lang_name = language_names.get(target_language.lower(), target_language)

        # 1. Collect all unique translatable text blocks
        text_blocks = set()

        def collect_from_container(container):
            for para in container.paragraphs:
                if para.text.strip():
                    text = para.text
                    # Map Wingdings characters before sending to LLM (Fix #1 from V3 report)
                    for char, mapped in self.wingdings_map.items():
                        text = text.replace(char, mapped)
                    text_blocks.add(text)
            
            # Recurse into tables which may contain more paragraphs
            tables = getattr(container, 'tables', [])
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        collect_from_container(cell)

        # Body content
        collect_from_container(doc)
        
        # Headers/Footers content (all sections and all types)
        for section in doc.sections:
            header_footer_objects = [
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer
            ]
            for obj in header_footer_objects:
                if obj: collect_from_container(obj)

        blocks = list(text_blocks)
        if not blocks:
            doc.save(output_path)
            return output_path

        # 2. Batch translation (parallel fan-out with per-batch fallback)
        batch_size = 20
        batches = [blocks[i:i+batch_size] for i in range(0, len(blocks), batch_size)]

        def _translate_batch(batch):
            try:
                protected_batch, batch_term_maps = [], []
                for seg in batch:
                    p, tm = self._protect_terms(seg)
                    protected_batch.append(p)
                    batch_term_maps.append(tm)

                prompt = f"""Translate each text segment to {target_lang_name}.

Return EXACTLY {len(batch)} lines — one translation per input segment.
- __TERM_N__ and __BRACKET_N__ tokens → copy verbatim, never translate.
- Checkboxes (☑ ☐ ✓ ✗ ■) → copy exactly as-is.
- No explanations, numbering, or markdown.

Segments:
---
{chr(10).join(protected_batch)}
---

Translated ({len(batch)} lines):"""
                response = self.client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": f"Professional corporate translator. Output exactly {len(batch)} lines matching the input count. Never add commentary or formatting. __TERM_N__ and __BRACKET_N__ tokens must be copied verbatim."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0
                )
                raw_output = response.choices[0].message.content.strip()
                if raw_output.startswith("```"):
                    lines = raw_output.split('\n')
                    if len(lines) > 2:
                        raw_output = '\n'.join(lines[1:-1]).strip()
                batch_translated_raw = [t.strip() for t in raw_output.split('\n') if t.strip()]
                batch_translated = [
                    self._restore_terms(batch_translated_raw[i], batch_term_maps[i])
                    if i < len(batch_translated_raw) else seg
                    for i, seg in enumerate(batch)
                ]
                return dict(zip(batch, batch_translated))
            except Exception as e:
                print(f"[OpenAI] Batch translation failed, keeping originals: {e}")
                return {seg: seg for seg in batch}

        # Run batches sequentially to stay within TPM limits
        translated_map = {}
        for batch in batches:
            translated_map.update(_translate_batch(batch))

        from docx.oxml.ns import qn

        def apply_to_container(container):
            is_hindi = target_language.lower() in ['hi', 'rajasthani']
            
            for para in container.paragraphs:
                # Need to map the para.text for key matching since we mapped text_blocks
                lookup_text = para.text
                for char, mapped in self.wingdings_map.items():
                    lookup_text = lookup_text.replace(char, mapped)
                
                if lookup_text in translated_map:
                    translated_text = translated_map[lookup_text]
                    if para.runs:
                        # PROBST preservation: Identify runs containing drawings or picts (images)
                        text_run_indices = []
                        for i, run in enumerate(para.runs):
                            has_drawing = run._element.findall('.//' + qn('w:drawing'))
                            has_pict = run._element.findall('.//' + qn('w:pict'))
                            if not has_drawing and not has_pict:
                                text_run_indices.append(i)
                        
                        if text_run_indices:
                            # Update the first text-only run and clear others
                            main_run = para.runs[text_run_indices[0]]
                            main_run.text = translated_text
                            
                            # Enforce Hindi font if applicable
                            if is_hindi:
                                main_run.font.name = 'Anek Devanagari'
                                # Explicitly set complex script font in XML
                                rPr = main_run._element.get_or_add_rPr()
                                rFonts = rPr.get_or_add_rFonts()
                                rFonts.set(qn('w:cs'), 'Anek Devanagari')
                                rFonts.set(qn('w:hAnsi'), 'Anek Devanagari')
                                rFonts.set(qn('w:ascii'), 'Anek Devanagari')

                            for idx in text_run_indices[1:]:
                                para.runs[idx].text = ""
                        else:
                            # Add a new run if all existing ones have drawings/picts
                            new_run = para.add_run(translated_text)
                            if is_hindi:
                                new_run.font.name = 'Anek Devanagari'
                                rPr = new_run._element.get_or_add_rPr()
                                rFonts = rPr.get_or_add_rFonts()
                                rFonts.set(qn('w:cs'), 'Anek Devanagari')
                    else:
                        new_run = para.add_run(translated_text)
                        if is_hindi:
                            new_run.font.name = 'Anek Devanagari'
                            rPr = new_run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:cs'), 'Anek Devanagari')
            
            tables = getattr(container, 'tables', [])
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        apply_to_container(cell)

        # Apply to all layers
        apply_to_container(doc)
        for section in doc.sections:
            for obj in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
                if obj: apply_to_container(obj)

        doc.save(output_path)
        return output_path
